"""
QP Extraction Module
Extracts question paper content from PDF and DOCX files and converts to structured format.
Auto-extracts mark distribution, CO mapping, Bloom's levels, and question text.
"""

import re
import io
from typing import List, Dict, Tuple, Optional
from pathlib import Path

try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from docx import Document
    from docx.table import Table
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class QPExtractionError(Exception):
    """Custom exception for QP extraction errors"""
    pass


class PDFQPExtractor:
    """Extracts question paper content from PDF files."""
    
    def __init__(self, file_obj):
        """
        Initialize with file object.
        
        Args:
            file_obj: File object or BytesIO containing PDF
        """
        if not HAS_PYPDF2:
            raise QPExtractionError("PyPDF2 not installed. Install via: pip install PyPDF2")
        
        self.file_obj = file_obj
        self.text = self._extract_text()
    
    def _extract_text(self) -> str:
        """Extract all text from PDF."""
        try:
            self.file_obj.seek(0)
            reader = PyPDF2.PdfReader(self.file_obj)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            raise QPExtractionError(f"Failed to extract PDF text: {str(e)}")
    
    def get_raw_text(self) -> str:
        """Return raw extracted text for debugging."""
        return self.text
    
    def extract_questions(self) -> Dict:
        """
        Extract questions from PDF text.
        
        Returns:
            {
                'part_a': [{'number': 1, 'text': '...', 'marks': 2}, ...],
                'part_b': [{'number': 11, 'option': 'a', 'text': '...', 'marks': 13}, ...],
                'part_c': [{'number': 16, 'text': '...', 'marks': 15}],
                'total_marks': 100,
                'errors': []
            }
        """
        try:
            part_a = self._extract_part_a()
            part_b = self._extract_part_b()
            part_c = self._extract_part_c()
            co_descriptions = self._extract_co_descriptions()
            
            # Extract CO/BL mapping from the PDF
            cobl_mapping = self._extract_cobl_mapping()
            
            # Apply CO/BL to the extracted questions
            part_a = self._apply_cobl_mapping(part_a, cobl_mapping)
            part_b = self._apply_cobl_mapping(part_b, cobl_mapping)
            part_c = self._apply_cobl_mapping(part_c, cobl_mapping)
            
            # Calculate total marks:
            # Part A: count of entries × 2
            # Part B: count of UNIQUE question numbers × 13 (each question has (a) OR (b) but counts as one)
            # Part C: 1 × 15 if exists
            
            part_a_marks = len(part_a) * 2
            
            # Get unique question numbers from Part B
            part_b_unique_questions = len(set(q['number'] for q in part_b)) if part_b else 0
            part_b_marks = part_b_unique_questions * 13
            
            part_c_marks = 15 if part_c else 0
            
            total_marks = part_a_marks + part_b_marks + part_c_marks
            
            errors = []
            if len(part_a) < 10:
                errors.append(f"Part A: Only extracted {len(part_a)}/10 questions")
            if part_b_unique_questions < 5:
                errors.append(f"Part B: Only extracted {part_b_unique_questions}/5 questions (with {len(part_b)} options total)")
            if not part_c:
                errors.append("Part C: Could not extract the compulsory question (Q16)")
            
            return {
                'part_a': part_a,
                'part_b': part_b,
                'part_c': part_c,
                'co_descriptions': co_descriptions,
                'total_marks': total_marks,
                'errors': errors
            }
        except Exception as e:
            return {
                'part_a': [],
                'part_b': [],
                'part_c': [],
                'co_descriptions': {},
                'total_marks': 0,
                'errors': [str(e)]
            }

    def _extract_co_descriptions(self) -> Dict[str, str]:
        """Extract CO1..CO5 description table/lines from PDF text."""
        co_descriptions: Dict[str, str] = {}

        # Common patterns in extracted PDF text:
        # CO1 <description>
        # CO2: <description>
        # C01 - <description>
        pattern = re.compile(
            r'\bC\s*O?\s*0?([1-5])\b\s*[:\-]?\s*(.+?)(?=(?:\n\s*\bC\s*O?\s*0?[1-5]\b)|(?:\n\s*BL\b)|(?:\n\s*PART\b)|$)',
            re.IGNORECASE | re.DOTALL,
        )

        for match in pattern.finditer(self.text):
            co_key = f"CO{match.group(1)}"
            desc = match.group(2).strip()
            desc = re.sub(r'\s+', ' ', desc)
            if desc and len(desc) > 3:
                co_descriptions[co_key] = desc

        # Line-by-line fallback for tightly formatted tables
        if len(co_descriptions) < 5:
            for line in self.text.splitlines():
                line = line.strip()
                if not line:
                    continue
                line_match = re.match(r'^C\s*O?\s*0?([1-5])\s*[:\-]?\s*(.+)$', line, re.IGNORECASE)
                if line_match:
                    co_key = f"CO{line_match.group(1)}"
                    desc = re.sub(r'\s+', ' ', line_match.group(2)).strip()
                    if desc and len(desc) > 3:
                        co_descriptions[co_key] = desc

        return co_descriptions
    
    def _extract_part(self, part: str) -> List[Dict]:
        """Extract questions for a specific part (A/B/C)."""
        questions = []
        
        if part == 'A':
            questions = self._extract_part_a()
        elif part == 'B':
            questions = self._extract_part_b()
        elif part == 'C':
            questions = self._extract_part_c()
        
        return questions
    
    def _extract_part_a(self) -> List[Dict]:
        """Extract Part A questions by simple text splitting on question numbers."""
        questions = []
        
        # Find Part A section between headers
        part_a_match = re.search(r'(?:PART-\s*A|Part\s+A)\s*(?:\(.*?\))?', self.text, re.IGNORECASE)
        if not part_a_match:
            return []
        
        start_idx = part_a_match.end()
        
        # Find where Part B starts
        part_b_match = re.search(r'(?:PART-\s*B|Part\s+B)\s*(?:\(.*?\))?', self.text[start_idx:], re.IGNORECASE)
        end_idx = start_idx + part_b_match.start() if part_b_match else len(self.text)
        
        part_a_text = self.text[start_idx:end_idx]
        
        # Find all question numbers (1-10) and their positions
        # Pattern: newline + optional spaces + question number
        for match in re.finditer(r'\n\s*([1-9]|10)(?:\s+|(?=[A-Z]))', part_a_text):
            q_num = int(match.group(1))
            text_start = match.end()
            
            # Find next question or end of section
            next_q_match = re.search(r'\n\s*([1-9]|10)(?:\s+|(?=[A-Z]))', part_a_text[text_start:])
            text_end = text_start + next_q_match.start() if next_q_match else len(part_a_text)
            
            q_text = part_a_text[text_start:text_end].strip()
            
            # Clean text: normalize whitespace, remove trailing marks
            q_text = re.sub(r'\n+', ' ', q_text)
            q_text = re.sub(r'\s+', ' ', q_text)
            q_text = re.sub(r'\s+\d+\s*$', '', q_text)  # Remove trailing mark number
            q_text = re.sub(r'\s+Page\s+\d+.*$', '', q_text, flags=re.IGNORECASE)
            q_text = q_text.strip()
            
            # Keep if meaningful
            if len(q_text) > 10:
                questions.append({
                    'number': q_num,
                    'text': q_text,
                    'marks': 2,
                    'option': None,
                    'co': 'CO1',
                    'bl': 'L1'
                })
        
        return questions

    
    def _extract_part_b(self) -> List[Dict]:
        """Extract Part B questions by simple text splitting on (a) and (b) markers."""
        questions = []
        
        # Find Part B section
        part_b_match = re.search(r'(?:PART-\s*B|Part\s+B)\s*(?:\(.*?\))?', self.text, re.IGNORECASE)
        if not part_b_match:
            return []
        
        start_idx = part_b_match.end()
        
        # Find where Part C starts
        part_c_match = re.search(r'(?:PART-\s*C|Part\s+C|PART-\s*0|PART-\s*G)', self.text[start_idx:], re.IGNORECASE)
        end_idx = start_idx + part_c_match.start() if part_c_match else len(self.text)
        
        part_b_text = self.text[start_idx:end_idx]
        
        # Split on (a) and (b) patterns
        # Look for question_number (a) or question_number (b)
        q_pattern = r'([1-5]|1[1-5]|W)\s*\(\s*([ab])\s*\)'
        
        matches = list(re.finditer(q_pattern, part_b_text, re.IGNORECASE))
        
        for i, match in enumerate(matches):
            q_num_str = match.group(1).strip()
            option = match.group(2).lower()
            
            # Parse question number
            if q_num_str.upper() == 'W':
                q_num = 15
            else:
                try:
                    q_num = int(q_num_str)
                    # If single digit, adjust to 11-15 range
                    if q_num < 11 and len(q_num_str) == 1:
                        q_num = 10 + q_num
                except ValueError:
                    continue
            
            # Verify range
            if q_num < 11 or q_num > 15:
                continue
            
            # Extract text from after this match until next (a)/(b) marker
            text_start = match.end()
            
            # Find next marker
            if i + 1 < len(matches):
                text_end = matches[i + 1].start()
            else:
                text_end = len(part_b_text)
            
            q_text = part_b_text[text_start:text_end].strip()
            
            # Clean text
            q_text = re.sub(r'\n+', ' ', q_text)
            q_text = re.sub(r'\s+', ' ', q_text)
            q_text = re.sub(r'\s+\d+\s*$', '', q_text)  # Remove trailing mark
            q_text = re.sub(r'\s+OR\s*$', '', q_text, flags=re.IGNORECASE)  # Remove OR
            q_text = re.sub(r'\s+Page\s+\d+.*$', '', q_text, flags=re.IGNORECASE)
            q_text = q_text.strip()
            
            # Keep if meaningful
            if len(q_text) >= 10:
                questions.append({
                    'number': q_num,
                    'text': q_text,
                    'marks': 13,
                    'option': option,
                    'co': 'CO1',
                    'bl': 'L1'
                })
        
        return sorted(questions, key=lambda x: (x['number'], x['option']))

    
    def _extract_part_c(self) -> List[Dict]:
        """Extract Part C question (1 compulsory question, 15 marks)."""
        # Find Part C section
        part_c_match = re.search(r'(?:PART-\s*[C0G]|Part\s+[C0G])\s*(?:\(.*?\))?', self.text, re.IGNORECASE)
        if not part_c_match:
            return []
        
        start_idx = part_c_match.end()
        part_c_text = self.text[start_idx:]
        
        # Look for question 16 marker
        q_match = re.search(r'(?:16\s*[.)]?|Q\.?No[.\s]*16)\s+(.+?)(?=\n\n|^Page|Answer|$)', part_c_text, re.IGNORECASE | re.DOTALL)
        
        if q_match:
            q_text = q_match.group(1).strip()
            
            # Clean text
            q_text = re.sub(r'\n+', ' ', q_text)
            q_text = re.sub(r'\s+', ' ', q_text)
            q_text = re.sub(r'\s+\d+\s*$', '', q_text)
            q_text = re.sub(r'\s+Page\s+\d+.*$', '', q_text, flags=re.IGNORECASE)
            q_text = q_text.strip()
            
            if len(q_text) > 15:
                return [{
                    'number': 16,
                    'text': q_text,
                    'marks': 15,
                    'option': None,
                    'co': 'CO1',
                    'bl': 'L1'
                }]
        
        return []

    
    def _extract_cobl_mapping(self) -> Dict[int, Dict[str, str]]:
        """Extract CO and BL values for each question from PDF.
        
        Handles table format: question_number | text | CO | BL
        or inline format: question_text CO1 L1
        """
        mapping = {}
        
        # Strategy: For each line, if it has a question number marker,
        # look for CO and BL values anywhere on that line
        
        lines = self.text.split('\n')
        
        for line in lines:
            if not line.strip():
                continue
            
            # Find ALL CO and BL patterns on this line
            co_matches = re.findall(r'CO([1-5])', line, re.IGNORECASE)
            bl_matches = re.findall(r'L([1-6])', line, re.IGNORECASE)
            
            if not co_matches or not bl_matches:
                continue
            
            # Get the first (or most relevant) CO and BL on this line
            co = f"CO{co_matches[0]}"
            bl = f"L{bl_matches[0]}"
            
            # Find question numbers on this line
            # Look for patterns like "1", "11", "11(a)", "W", etc.
            q_patterns = [
                r'\b([1-9]|1[0-6])\b(?:\s*\(|[^0-9]|$)',  # Single/double digit question
                r'\b(W)\b',  # OCR error for 15
            ]
            
            for pattern in q_patterns:
                q_matches = re.findall(pattern, line)
                for q_match in q_matches:
                    try:
                        if q_match.upper() == 'W':
                            q_num = 15
                        else:
                            q_num = int(q_match)
                        
                        # Valid question range
                        if 1 <= q_num <= 16:
                            mapping[q_num] = {'co': co, 'bl': bl}
                    except ValueError:
                        pass
        
        return mapping
    
    def _apply_cobl_mapping(self, questions: List[Dict], mapping: Dict) -> List[Dict]:
        """Apply CO/BL mapping to extracted questions."""
        for q in questions:
            q_num = q['number']
            if q_num in mapping:
                q['co'] = mapping[q_num].get('co', q['co'])
                q['bl'] = mapping[q_num].get('bl', q['bl'])
        return questions

    
    def _parse_question_number(self, prefix: str, part: str) -> Optional[int]:
        """Parse question number from match prefix."""
        match = re.search(r'\d+', prefix)
        if match:
            return int(match.group(0))
        return None


class DOCXQPExtractor:
    """Extracts question paper content from DOCX files."""
    
    def __init__(self, file_obj):
        """
        Initialize with file object.
        
        Args:
            file_obj: File object or BytesIO containing DOCX
        """
        if not HAS_DOCX:
            raise QPExtractionError("python-docx not installed. Install via: pip install python-docx")
        
        self.file_obj = file_obj
        self.doc = self._load_document()
    
    def _load_document(self):
        """Load DOCX document."""
        try:
            self.file_obj.seek(0)
            return Document(self.file_obj)
        except Exception as e:
            raise QPExtractionError(f"Failed to load DOCX: {str(e)}")
    
    def get_raw_text(self) -> str:
        """Return all document text for debugging."""
        return '\n'.join(paragraph.text for paragraph in self.doc.paragraphs)

    def _extract_cobl_from_text(self, text: str) -> Tuple[Optional[str], Optional[str]]:
        """Extract CO and Bloom's level from text snippets like CO3, L4."""
        co = self._normalize_co_value(text)
        bl = self._normalize_bl_value(text)
        return co, bl

    def _normalize_co_value(self, text: str) -> Optional[str]:
        """Normalize CO formats like CO3, CO-3, C03, or plain 3 to CO3."""
        if not text:
            return None

        co_match = re.search(r'\bCO\s*[-:]?\s*([1-5])\b', text, re.IGNORECASE)
        if not co_match:
            co_match = re.search(r'\bC0*([1-5])\b', text, re.IGNORECASE)
        if not co_match:
            stripped = text.strip()
            if re.fullmatch(r'[1-5]', stripped):
                co_match = re.match(r'([1-5])', stripped)

        return f"CO{co_match.group(1)}" if co_match else None

    def _normalize_bl_value(self, text: str) -> Optional[str]:
        """Normalize BL formats like L3, BL3, Bloom 3, or plain 3 to L3."""
        if not text:
            return None

        bl_match = re.search(r'\bL\s*[-:]?\s*([1-6])\b', text, re.IGNORECASE)
        if not bl_match:
            bl_match = re.search(r'\bBL\s*[-:]?\s*([1-6])\b', text, re.IGNORECASE)
        if not bl_match:
            bl_match = re.search(r'\bBLOOM\w*\s*[-:]?\s*([1-6])\b', text, re.IGNORECASE)
        if not bl_match:
            stripped = text.strip()
            if re.fullmatch(r'[1-6]', stripped):
                bl_match = re.match(r'([1-6])', stripped)

        return f"L{bl_match.group(1)}" if bl_match else None

    def _detect_table_header_indices(self, cells: List[str]) -> Dict[str, Optional[int]]:
        """Detect common QP table header columns."""
        indices = {
            'qno': None,
            'question': None,
            'option': None,
            'co': None,
            'bl': None,
        }

        for idx, cell in enumerate(cells):
            header = cell.strip().lower()
            if not header:
                continue

            if re.search(r'question\s*no|q\.?\s*no|^s\.?\s*no|^no\.?$', header):
                indices['qno'] = idx
            elif 'question' in header and 'text' in header:
                indices['question'] = idx
            elif header == 'question' and indices['question'] is None:
                indices['question'] = idx
            elif 'option' in header:
                indices['option'] = idx
            elif re.search(r'course\s*outcome|\bco\b', header):
                indices['co'] = idx
            elif re.search(r'bloom|\bbl\b', header):
                indices['bl'] = idx

        return indices

    def _is_header_row(self, cells: List[str]) -> bool:
        """Check if a row looks like a table header row."""
        row_text = ' '.join(cells).lower()
        return any(
            key in row_text
            for key in ['question no', 'q no', 'course outcome', 'bloom', 'marks', 'option']
        )

    def _extract_option_value(self, text: str) -> Optional[str]:
        """Extract option marker a/b from text."""
        match = re.search(r'\(([ab])\)', text, re.IGNORECASE)
        if not match:
            match = re.search(r'\b([ab])\b', text, re.IGNORECASE)
        return match.group(1).lower() if match else None

    def _clean_question_text(self, text: str) -> str:
        """Remove inline CO/BL tokens from question text for cleaner display."""
        cleaned = re.sub(r'\bCO\s*[1-5]\b', '', text, flags=re.IGNORECASE)
        cleaned = re.sub(r'\bL\s*[1-6]\b', '', cleaned, flags=re.IGNORECASE)

        # Remove trailing marks-column values that may leak into question text.
        # Common structured QP mark values are 2, 13, and 15.
        cleaned = re.sub(r'\s+(?:\(?\s*(2|13|15)\s*\)?|\[\s*(2|13|15)\s*\])\s*$', '', cleaned)

        cleaned = re.sub(r'\s+', ' ', cleaned)
        return cleaned.strip()
    
    def extract_questions(self) -> Dict:
        """
        Extract questions from DOCX.
        
        Returns:
            {
                'part_a': [{'number': 1, 'text': '...', 'marks': 2}, ...],
                'part_b': [{'number': 11, 'option': 'a', 'text': '...', 'marks': 13}, ...],
                'part_c': [{'number': 16, 'text': '...', 'marks': 15}],
                'total_marks': 100,
                'errors': []
            }
        """
        try:
            co_descriptions = self._extract_co_descriptions()

            # Try to extract from tables first (common QP format)
            table_data = self._extract_from_tables()
            if table_data['part_a'] or table_data['part_b']:
                table_data['co_descriptions'] = co_descriptions
                return table_data
            
            # Fallback to paragraph extraction
            paragraph_data = self._extract_from_paragraphs()
            paragraph_data['co_descriptions'] = co_descriptions
            return paragraph_data
        
        except Exception as e:
            return {
                'part_a': [],
                'part_b': [],
                'part_c': [],
                'co_descriptions': {},
                'total_marks': 0,
                'errors': [str(e)]
            }

    def _extract_co_descriptions(self) -> Dict[str, str]:
        """Extract CO1..CO5 descriptions from DOCX tables/paragraphs."""
        co_descriptions: Dict[str, str] = {}

        # Preferred: dedicated CO table rows
        for table in self.doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if not any(cells):
                    continue

                row_text = ' '.join(cells)
                co_match = re.search(r'\bC\s*O?\s*0?([1-5])\b', row_text, re.IGNORECASE)
                if not co_match:
                    continue

                co_key = f"CO{co_match.group(1)}"

                # Try to pick description from non-CO cell first.
                desc_candidates = []
                for cell in cells:
                    cleaned_cell = re.sub(r'\s+', ' ', cell).strip()
                    if not cleaned_cell:
                        continue
                    if re.fullmatch(r'C\s*O?\s*0?[1-5]', cleaned_cell, re.IGNORECASE):
                        continue
                    if re.search(r'\bC\s*O?\s*0?[1-5]\b', cleaned_cell, re.IGNORECASE):
                        cleaned_cell = re.sub(r'\bC\s*O?\s*0?[1-5]\b\s*[:\-]?', '', cleaned_cell, flags=re.IGNORECASE).strip()
                    if cleaned_cell:
                        desc_candidates.append(cleaned_cell)

                if desc_candidates:
                    co_descriptions[co_key] = max(desc_candidates, key=len)

        # Fallback: paragraph lines like "CO1: ..."
        if len(co_descriptions) < 5:
            for paragraph in self.doc.paragraphs:
                line = paragraph.text.strip()
                if not line:
                    continue
                line_match = re.match(r'^C\s*O?\s*0?([1-5])\s*[:\-]?\s*(.+)$', line, re.IGNORECASE)
                if line_match:
                    co_key = f"CO{line_match.group(1)}"
                    desc = re.sub(r'\s+', ' ', line_match.group(2)).strip()
                    if desc and len(desc) > 3:
                        co_descriptions[co_key] = desc

        return co_descriptions
    
    def _extract_from_tables(self) -> Dict:
        """Extract questions from DOCX tables (preferred method)."""
        part_a = []
        part_b = []
        part_c = []
        seen_a = set()
        seen_b = set()
        part_b_option_counter = {}
        
        for table in self.doc.tables:
            header_indices = {
                'qno': None,
                'question': None,
                'option': None,
                'co': None,
                'bl': None,
            }

            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if not any(cells):
                    continue

                if self._is_header_row(cells):
                    header_indices = self._detect_table_header_indices(cells)
                    continue
                
                joined_row = ' '.join(cells)

                q_num = None
                if header_indices['qno'] is not None and header_indices['qno'] < len(cells):
                    q_num = self._extract_number(cells[header_indices['qno']])
                if q_num is None:
                    q_num = self._extract_number(cells[0]) if cells else None
                if q_num is None:
                    q_num = self._extract_number(joined_row)

                if q_num is None or not (1 <= q_num <= 16):
                    continue

                co_val = None
                bl_val = None
                if header_indices['co'] is not None and header_indices['co'] < len(cells):
                    co_val = self._normalize_co_value(cells[header_indices['co']])
                if header_indices['bl'] is not None and header_indices['bl'] < len(cells):
                    bl_val = self._normalize_bl_value(cells[header_indices['bl']])

                inline_co, inline_bl = self._extract_cobl_from_text(joined_row)
                co_val = co_val or inline_co or 'CO1'
                bl_val = bl_val or inline_bl or 'L1'

                q_text = None
                if header_indices['question'] is not None and header_indices['question'] < len(cells):
                    q_text = cells[header_indices['question']]
                if not q_text:
                    excluded_indices = {
                        header_indices['qno'],
                        header_indices['co'],
                        header_indices['bl'],
                        header_indices['option'],
                    }
                    text_parts = [
                        value for idx, value in enumerate(cells)
                        if idx not in excluded_indices and value
                    ]
                    q_text = ' '.join(text_parts) if text_parts else joined_row

                q_text = self._clean_question_text(q_text)
                if len(q_text) < 10:
                    continue

                if 1 <= q_num <= 10:
                    if q_num in seen_a:
                        continue
                    seen_a.add(q_num)
                    part_a.append({
                        'number': q_num,
                        'text': q_text,
                        'marks': 2,
                        'option': None,
                        'co': co_val,
                        'bl': bl_val
                    })
                elif 11 <= q_num <= 15:
                    option = None
                    if header_indices['option'] is not None and header_indices['option'] < len(cells):
                        option = self._extract_option_value(cells[header_indices['option']])
                    if not option:
                        option = self._extract_option_value(joined_row)
                    if not option:
                        used = part_b_option_counter.get(q_num, 0)
                        option = 'a' if used % 2 == 0 else 'b'
                    part_b_option_counter[q_num] = part_b_option_counter.get(q_num, 0) + 1

                    key = (q_num, option)
                    if key in seen_b:
                        continue
                    seen_b.add(key)
                    part_b.append({
                        'number': q_num,
                        'text': q_text,
                        'marks': 13,
                        'option': option,
                        'co': co_val,
                        'bl': bl_val
                    })
                else:
                    if part_c:
                        continue
                    part_c.append({
                        'number': 16,
                        'text': q_text,
                        'marks': 15,
                        'option': None,
                        'co': co_val,
                        'bl': bl_val
                    })

        part_a = sorted(part_a, key=lambda x: x['number'])
        part_b = sorted(part_b, key=lambda x: (x['number'], x['option']))
        part_b_unique_questions = len(set(q['number'] for q in part_b)) if part_b else 0
        total_marks = len(part_a) * 2 + part_b_unique_questions * 13 + (15 if part_c else 0)
        
        return {
            'part_a': part_a[:10],
            'part_b': part_b[:10],
            'part_c': part_c[:1],
            'total_marks': total_marks,
            'errors': []
        }
    
    def _extract_from_paragraphs(self) -> Dict:
        """Fallback: Extract from paragraphs if table parsing fails."""
        text = self.get_raw_text()
        
        # Split by Part markers
        part_a_match = re.search(r'Part\s*A(.*?)(?=Part\s*B|$)', text, re.IGNORECASE | re.DOTALL)
        part_b_match = re.search(r'Part\s*B(.*?)(?=Part\s*C|$)', text, re.IGNORECASE | re.DOTALL)
        part_c_match = re.search(r'Part\s*C(.*?)$', text, re.IGNORECASE | re.DOTALL)
        
        part_a = []
        part_b = []
        part_c = []
        
        # Extract Part A
        if part_a_match:
            part_a_text = part_a_match.group(1)
            questions = re.split(r'\n(?=\s*(?:[1-9]|10)\s*(?:[.)]|\s))', part_a_text)
            for idx, q in enumerate(questions[:10]):
                q = q.strip()
                if q:
                    q_num = self._extract_number(q) or idx + 1
                    co_val, bl_val = self._extract_cobl_from_text(q)
                    part_a.append({
                        'number': q_num,
                        'text': self._clean_question_text(q),
                        'marks': 2,
                        'option': None,
                        'co': co_val or 'CO1',
                        'bl': bl_val or 'L1'
                    })
        
        # Extract Part B
        if part_b_match:
            part_b_text = part_b_match.group(1)
            questions = re.split(r'\n(?=\s*(?:1[1-5])\s*(?:[.)]|\(|\s)|\s*\([ab]\))', part_b_text)
            for idx, q in enumerate(questions[:10]):
                q = q.strip()
                if q:
                    q_num = self._extract_number(q) or (11 + idx // 2)
                    co_val, bl_val = self._extract_cobl_from_text(q)
                    option_match = re.search(r'\(([ab])\)', q, re.IGNORECASE)
                    option = option_match.group(1).lower() if option_match else ('a' if idx % 2 == 0 else 'b')
                    part_b.append({
                        'number': q_num,
                        'text': self._clean_question_text(q),
                        'marks': 13,
                        'option': option,
                        'co': co_val or 'CO1',
                        'bl': bl_val or 'L1'
                    })
        
        # Extract Part C
        if part_c_match:
            part_c_text = part_c_match.group(1).strip()
            if part_c_text:
                co_val, bl_val = self._extract_cobl_from_text(part_c_text)
                part_c.append({
                    'number': 16,
                    'text': self._clean_question_text(part_c_text),
                    'marks': 15,
                    'option': None,
                    'co': co_val or 'CO1',
                    'bl': bl_val or 'L1'
                })

        part_b_unique_questions = len(set(q['number'] for q in part_b)) if part_b else 0
        total_marks = len(part_a) * 2 + part_b_unique_questions * 13 + (15 if part_c else 0)
        
        return {
            'part_a': part_a,
            'part_b': part_b,
            'part_c': part_c,
            'total_marks': total_marks,
            'errors': []
        }
    
    def _extract_number(self, text: str) -> Optional[int]:
        """Extract question number from text."""
        match = re.search(r'^(\d+)', text.strip())
        if match:
            return int(match.group(1))
        return None


def extract_qp_from_file(file_obj, filename: str) -> Dict:
    """
    Universal extractor: routes to PDF or DOCX extractor based on file extension.
    
    Args:
        file_obj: File object
        filename: Original filename to determine type
    
    Returns:
        Extraction result dict with parts and errors
    """
    filename_lower = filename.lower()
    
    try:
        if filename_lower.endswith('.pdf'):
            if not HAS_PYPDF2:
                return {
                    'part_a': [], 'part_b': [], 'part_c': [],
                    'co_descriptions': {},
                    'total_marks': 0,
                    'errors': ['PDF support not available. Install: pip install PyPDF2']
                }
            extractor = PDFQPExtractor(file_obj)
        
        elif filename_lower.endswith('.docx'):
            if not HAS_DOCX:
                return {
                    'part_a': [], 'part_b': [], 'part_c': [],
                    'co_descriptions': {},
                    'total_marks': 0,
                    'errors': ['DOCX support not available. Install: pip install python-docx']
                }
            extractor = DOCXQPExtractor(file_obj)
        
        else:
            return {
                'part_a': [], 'part_b': [], 'part_c': [],
                'co_descriptions': {},
                'total_marks': 0,
                'errors': [f'Unsupported file type: {filename}. Only PDF and DOCX are supported.']
            }
        
        result = extractor.extract_questions()

        # Helpful hint: PDF table extraction is often lossy for CO/BL columns.
        if filename_lower.endswith('.pdf'):
            extracted_questions = result.get('part_a', []) + result.get('part_b', []) + result.get('part_c', [])
            has_non_default_cobl = any(
                (q.get('co') and q.get('co') != 'CO1') or (q.get('bl') and q.get('bl') != 'L1')
                for q in extracted_questions
            )
            if extracted_questions and not has_non_default_cobl:
                result.setdefault('errors', []).append(
                    "CO/BL mapping could not be reliably read from this PDF. For better CO/BL extraction, upload the same paper as DOCX."
                )

        return result
    
    except QPExtractionError as e:
        return {
            'part_a': [], 'part_b': [], 'part_c': [],
            'co_descriptions': {},
            'total_marks': 0,
            'errors': [str(e)]
        }
    except Exception as e:
        return {
            'part_a': [], 'part_b': [], 'part_c': [],
            'co_descriptions': {},
            'total_marks': 0,
            'errors': [f'Unexpected error during extraction: {str(e)}']
        }


def create_qp_questions_from_extraction(qp_instance, extraction_result: Dict, bloom_mapping: Dict = None, co_mapping: Dict = None) -> Tuple[int, List[str]]:
    """
    Create QPQuestion instances from extraction result.
    
    Args:
        qp_instance: StructuredQuestionPaper instance (must be saved)
        extraction_result: Result from extract_qp_from_file()
        bloom_mapping: Optional dict mapping question indices to Bloom levels
        co_mapping: Optional dict mapping question indices to Course Outcomes
    
    Returns:
        (questions_created_count, errors_list)
    """
    from main_app.models import QPQuestion
    
    questions_created = 0
    errors = []
    
    if not qp_instance.id:
        errors.append("QP instance must be saved before creating questions")
        return 0, errors
    
    try:
        # Default CO/Bloom if not provided
        if bloom_mapping is None:
            bloom_mapping = {}
        if co_mapping is None:
            co_mapping = {}
        
        # Part A
        for idx, q in enumerate(extraction_result.get('part_a', [])):
            try:
                question = QPQuestion.objects.create(
                    question_paper=qp_instance,
                    part='A',
                    question_number=q.get('number', idx + 1),
                    question_text=q.get('text', ''),
                    marks=q.get('marks', 2),
                    course_outcome=co_mapping.get(f'a_{idx}', 'CO1'),
                    bloom_level=bloom_mapping.get(f'a_{idx}', 'L1'),
                    is_or_option=False,
                    or_pair_number=None,
                    option_label='',  # Empty string instead of None for non-option questions
                    has_subdivisions=False
                )
                questions_created += 1
            except Exception as e:
                errors.append(f"Part A Q{idx + 1}: {str(e)}")
        
        # Part B
        for idx, q in enumerate(extraction_result.get('part_b', [])):
            try:
                or_pair = 11 + (idx // 2)
                option_label = '(a)' if idx % 2 == 0 else '(b)'
                
                question = QPQuestion.objects.create(
                    question_paper=qp_instance,
                    part='B',
                    question_number=or_pair,
                    question_text=q.get('text', ''),
                    marks=q.get('marks', 13),
                    course_outcome=co_mapping.get(f'b_{idx}', 'CO1'),
                    bloom_level=bloom_mapping.get(f'b_{idx}', 'L1'),
                    is_or_option=True,
                    or_pair_number=or_pair,
                    option_label=option_label,
                    has_subdivisions=False
                )
                questions_created += 1
            except Exception as e:
                errors.append(f"Part B Q{11 + idx // 2}{option_label}: {str(e)}")
        
        # Part C
        for idx, q in enumerate(extraction_result.get('part_c', [])):
            try:
                question = QPQuestion.objects.create(
                    question_paper=qp_instance,
                    part='C',
                    question_number=16,
                    question_text=q.get('text', ''),
                    marks=q.get('marks', 15),
                    course_outcome=co_mapping.get('c_0', 'CO1'),
                    bloom_level=bloom_mapping.get('c_0', 'L1'),
                    is_or_option=False,
                    or_pair_number=None,
                    option_label='',  # Empty string instead of None for non-option questions
                    has_subdivisions=False
                )
                questions_created += 1
            except Exception as e:
                errors.append(f"Part C: {str(e)}")
    
    except Exception as e:
        errors.append(f"Unexpected error creating questions: {str(e)}")
    
    return questions_created, errors
