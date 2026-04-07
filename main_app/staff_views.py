"""
Anna University CSE Department ERP System
Faculty (Staff) Views
"""

import json
import uuid
from datetime import date, datetime
from django.utils import timezone

from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.core.files.storage import FileSystemStorage
from django.http import HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.views.decorators.csrf import csrf_exempt
from django.db.models import Count, Q
from collections import defaultdict

from .forms import (
    LeaveRequestForm, FeedbackForm, PublicationForm, 
    FacultyProfileEditForm, AccountUserForm, BulkAttendanceForm,
    QuestionPaperSubmissionForm
)
from .models import (
    Account_User, Faculty_Profile, Student_Profile, Course_Assignment,
    Attendance, LeaveRequest, Feedback, Notification, Publication,
    Announcement, AcademicYear, Semester, QuestionPaperAssignment,
    Timetable, TimetableEntry, TimeSlot, StructuredQuestionPaper, QPQuestion,
    ProgramBatch, RegulationCoursePlan, ElectiveOfferingFacultyAssignment
)
from .utils.web_scrapper import fetch_acoe_updates
from .utils.cir_scrapper import fetch_cir_ticker_announcements
from .utils.qp_checklist_doc import append_checklist_to_document

def check_faculty_permission(user, request=None):
    """
    Check if user is Faculty, Guest Faculty, or HOD in faculty mode.
    HOD is identified via Faculty_Profile.designation == 'HOD'.
    """
    if not user.is_authenticated:
        return False
    
    # Regular faculty or guest
    if user.role in ['FACULTY', 'GUEST']:
        return True
    
    # HOD (via Faculty_Profile.designation) can access faculty views
    if user.is_hod:
        return True
    
    return False


def _ensure_faculty_profile(user):
    """Create a minimal faculty profile when role is faculty-like but profile is missing."""
    faculty = Faculty_Profile.objects.filter(user=user).first()
    if faculty:
        return faculty, False

    staff_id = f"TEMP_{uuid.uuid4().hex[:8].upper()}"
    while Faculty_Profile.objects.filter(staff_id=staff_id).exists():
        staff_id = f"TEMP_{uuid.uuid4().hex[:8].upper()}"

    faculty = Faculty_Profile.objects.create(
        user=user,
        staff_id=staff_id,
        designation='AP',
        is_external=(user.role == 'GUEST'),
    )
    return faculty, True


def _build_auto_distribution_checklist(qp):
    """Build auto-filled checklist Form 2 data from stored QP questions."""
    co_keys = ['CO1', 'CO2', 'CO3', 'CO4', 'CO5']
    bl_keys = ['L1', 'L2', 'L3', 'L4', 'L5', 'L6']

    part_a_map = {q.question_number: q for q in qp.questions.filter(part='A')}

    part_b_map = {}
    for q_num in range(11, 16):
        pair_questions = list(qp.questions.filter(part='B', or_pair_number=q_num).order_by('option_label', 'id'))
        selected = next((q for q in pair_questions if q.option_label == '(a)'), None)
        if not selected and pair_questions:
            selected = pair_questions[0]
        if selected:
            part_b_map[q_num] = selected

    part_c_q = qp.questions.filter(part='C').order_by('id').first()

    rows = []
    co_totals = {key: 0 for key in co_keys}
    bl_totals = {key: 0 for key in bl_keys}

    def add_row(question_no, q_obj, expected_marks):
        marks = q_obj.marks if q_obj else expected_marks
        co = (q_obj.course_outcome or '').strip() if q_obj else ''
        bl = (q_obj.bloom_level or '').strip() if q_obj else ''

        co_marks = {key: (marks if co == key else 0) for key in co_keys}
        bl_marks = {key: (marks if bl == key else 0) for key in bl_keys}

        for key in co_keys:
            co_totals[key] += co_marks[key]
        for key in bl_keys:
            bl_totals[key] += bl_marks[key]

        rows.append({
            'question_no': question_no,
            'co_marks': co_marks,
            'total_marks': marks,
            'bl_marks': bl_marks,
        })

    for q_num in range(1, 11):
        add_row(q_num, part_a_map.get(q_num), 2)

    for q_num in range(11, 16):
        add_row(q_num, part_b_map.get(q_num), 13)

    add_row(16, part_c_q, 15)

    total_marks = sum(r['total_marks'] for r in rows)
    l1_l2_total = bl_totals['L1'] + bl_totals['L2']
    l3_l4_total = bl_totals['L3'] + bl_totals['L4']
    l5_l6_total = bl_totals['L5'] + bl_totals['L6']

    return {
        'rows': rows,
        'co_totals': co_totals,
        'bl_totals': bl_totals,
        'total_marks': total_marks,
        'l1_l2_total': l1_l2_total,
        'l3_l4_total': l3_l4_total,
        'l5_l6_total': l5_l6_total,
        'l1_l2_percentage': round((l1_l2_total / total_marks * 100), 2) if total_marks else 0,
        'l3_l4_percentage': round((l3_l4_total / total_marks * 100), 2) if total_marks else 0,
        'l5_l6_percentage': round((l5_l6_total / total_marks * 100), 2) if total_marks else 0,
    }


# =============================================================================
# DASHBOARD
# =============================================================================

@login_required
def staff_home(request):
    """Faculty Dashboard"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied. Faculty privileges required.")
        return redirect('/')
    
    faculty, profile_created = _ensure_faculty_profile(request.user)
    if profile_created:
        messages.warning(
            request,
            "Your faculty profile was missing and has been recreated automatically. Please verify profile details.",
        )
    
    # Get assigned courses
    assignments = Course_Assignment.objects.filter(faculty=faculty, is_active=True)
    total_courses = assignments.count()
    
    # Count students across all assigned batches
    total_students = 0
    for assignment in assignments:
        student_count = Student_Profile.objects.filter(batch_label=assignment.batch_label).count()
        total_students += student_count
    
    # Leave and attendance stats
    total_leaves = LeaveRequest.objects.filter(user=request.user).count()
    pending_leaves = LeaveRequest.objects.filter(user=request.user, status='PENDING').count()
    
    # Attendance stats per course
    course_list = []
    attendance_count_list = []
    for assignment in assignments:
        attendance_count = Attendance.objects.filter(assignment=assignment).count()
        course_list.append(assignment.course.title[:15])
        attendance_count_list.append(attendance_count)
    
    # Publications
    publications = Publication.objects.filter(faculty=faculty)
    total_publications = publications.count()
    unverified_publications = publications.filter(is_verified=False).count()
    
    # Notifications
    notifications = Notification.objects.filter(recipient=request.user, is_read=False)[:5]
    
    # Fetch announcements
    announcements = []
    try:
        acoe_updates = fetch_acoe_updates()
        announcements.extend(acoe_updates)
    except:
        pass
    
    try:
        cir_announcements = fetch_cir_ticker_announcements(limit=5)
        announcements.extend(cir_announcements)
    except:
        pass
    
    # Department announcements
    dept_announcements = Announcement.objects.filter(
        is_active=True, 
        audience__in=['ALL', 'FACULTY']
    )[:5]
    
    context = {
        'page_title': f'Faculty Dashboard - {faculty.user.full_name}',
        'faculty': faculty,
        'total_courses': total_courses,
        'total_students': total_students,
        'total_leaves': total_leaves,
        'pending_leaves': pending_leaves,
        'course_list': json.dumps(course_list),
        'attendance_count_list': json.dumps(attendance_count_list),
        'total_publications': total_publications,
        'unverified_publications': unverified_publications,
        'notifications': notifications,
        'announcements': announcements,
        'dept_announcements': dept_announcements,
        'assignments': assignments,
    }
    return render(request, 'staff_template/home_content.html', context)


# =============================================================================
# ATTENDANCE MANAGEMENT
# =============================================================================

@login_required
def staff_take_attendance(request):
    """Take attendance for assigned courses"""
    if not check_faculty_permission(request.user):
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    assignments = Course_Assignment.objects.filter(faculty=faculty, is_active=True)
    
    context = {
        'assignments': assignments,
        'page_title': 'Take Attendance'
    }
    return render(request, 'staff_template/staff_take_attendance.html', context)


@csrf_exempt
@login_required
def get_students(request):
    """Get students for a course assignment"""
    if not check_faculty_permission(request.user):
        return JsonResponse({'error': 'Permission denied'}, status=403)
    
    assignment_id = request.POST.get('assignment')
    try:
        assignment = get_object_or_404(Course_Assignment, id=assignment_id)
        
        # Get students matching the batch label
        students = Student_Profile.objects.filter(batch_label=assignment.batch_label).select_related('user')
        
        student_data = []
        for student in students:
            data = {
                "id": str(student.id),
                "name": student.user.full_name,
                "register_no": student.register_no
            }
            student_data.append(data)
        
        return JsonResponse(json.dumps(student_data), content_type='application/json', safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@csrf_exempt
@login_required
def save_attendance(request):
    """Save attendance records"""
    if not check_faculty_permission(request.user):
        return HttpResponse("Permission denied", status=403)
    
    student_data = request.POST.get('student_ids')
    attendance_date = request.POST.get('date')
    assignment_id = request.POST.get('assignment')
    period = request.POST.get('period', 1)
    
    students = json.loads(student_data)
    
    try:
        assignment = get_object_or_404(Course_Assignment, id=assignment_id)
        
        for student_dict in students:
            student = get_object_or_404(Student_Profile, id=student_dict.get('id'))
            
            # Check if attendance already exists
            attendance, created = Attendance.objects.update_or_create(
                student=student,
                assignment=assignment,
                date=attendance_date,
                period=period,
                defaults={
                    'status': 'PRESENT' if student_dict.get('status') else 'ABSENT'
                }
            )
        
        return HttpResponse("OK")
    except Exception as e:
        return HttpResponse(f"Error: {str(e)}", status=400)


@login_required
def staff_update_attendance(request):
    """Update attendance records"""
    if not check_faculty_permission(request.user):
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    assignments = Course_Assignment.objects.filter(faculty=faculty, is_active=True)
    
    context = {
        'assignments': assignments,
        'page_title': 'Update Attendance'
    }
    return render(request, 'staff_template/staff_update_attendance.html', context)


@csrf_exempt
@login_required
def get_student_attendance(request):
    """Get attendance data for editing"""
    if not check_faculty_permission(request.user):
        return JsonResponse({'error': 'Permission denied'}, status=403)
    
    assignment_id = request.POST.get('assignment')
    attendance_date_id = request.POST.get('attendance_date_id')
    period = request.POST.get('period', 1)
    
    try:
        assignment = get_object_or_404(Course_Assignment, id=assignment_id)
        
        # Get the date from attendance record
        attendance_record = Attendance.objects.filter(id=attendance_date_id).first()
        if attendance_record:
            attendance_date = attendance_record.date
        else:
            return JsonResponse(json.dumps([]), content_type='application/json', safe=False)
        
        # Get students for this batch
        students = Student_Profile.objects.filter(batch_label=assignment.batch_label).select_related('user')
        
        student_data = []
        for student in students:
            # Check if attendance exists
            try:
                attendance = Attendance.objects.get(
                    student=student, 
                    assignment=assignment,
                    date=attendance_date,
                    period=period
                )
                status = attendance.status == 'PRESENT'
            except Attendance.DoesNotExist:
                status = False
            
            data = {
                "id": str(student.id),
                "name": student.user.full_name,
                "register_no": student.register_no,
                "status": status
            }
            student_data.append(data)
        
        return JsonResponse(json.dumps(student_data), content_type='application/json', safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@csrf_exempt
@login_required
def update_attendance(request):
    """Update existing attendance records"""
    if not check_faculty_permission(request.user):
        return HttpResponse("Permission denied", status=403)
    
    student_data = request.POST.get('student_ids')
    attendance_date_id = request.POST.get('date')
    assignment_id = request.POST.get('assignment')
    period = request.POST.get('period', 1)
    
    students = json.loads(student_data)
    
    try:
        assignment = get_object_or_404(Course_Assignment, id=assignment_id)
        
        # Get the date from attendance record
        attendance_record = Attendance.objects.filter(id=attendance_date_id).first()
        if attendance_record:
            attendance_date = attendance_record.date
        else:
            return HttpResponse("Invalid attendance record", status=400)
        
        for student_dict in students:
            student = get_object_or_404(Student_Profile, id=student_dict.get('id'))
            
            Attendance.objects.update_or_create(
                student=student,
                assignment=assignment,
                date=attendance_date,
                period=period,
                defaults={
                    'status': 'PRESENT' if student_dict.get('status') else 'ABSENT'
                }
            )
        
        return HttpResponse("OK")
    except Exception as e:
        return HttpResponse(f"Error: {str(e)}", status=400)


# =============================================================================
# LEAVE MANAGEMENT
# =============================================================================

@login_required
def staff_apply_leave(request):
    """Apply for leave"""
    if not check_faculty_permission(request.user):
        return redirect('/')
    
    form = LeaveRequestForm(request.POST or None, request.FILES or None)
    leave_history = LeaveRequest.objects.filter(user=request.user).order_by('-created_at')
    
    context = {
        'form': form,
        'leave_history': leave_history,
        'page_title': 'Apply for Leave'
    }
    
    if request.method == 'POST':
        if form.is_valid():
            try:
                leave = form.save(commit=False)
                leave.user = request.user
                leave.save()
                messages.success(request, "Leave application submitted for review")
                return redirect(reverse('staff_apply_leave'))
            except Exception as e:
                messages.error(request, f"Could not apply: {str(e)}")
        else:
            messages.error(request, "Please correct the errors in the form")
    
    return render(request, "staff_template/staff_apply_leave.html", context)


# =============================================================================
# FEEDBACK
# =============================================================================

@login_required
def staff_feedback(request):
    """Submit feedback"""
    if not check_faculty_permission(request.user):
        return redirect('/')
    
    form = FeedbackForm(request.POST or None)
    feedbacks = Feedback.objects.filter(user=request.user).order_by('-created_at')
    
    context = {
        'form': form,
        'feedbacks': feedbacks,
        'page_title': 'Submit Feedback'
    }
    
    if request.method == 'POST':
        if form.is_valid():
            try:
                feedback = form.save(commit=False)
                feedback.user = request.user
                feedback.save()
                messages.success(request, "Feedback submitted successfully")
                return redirect(reverse('staff_feedback'))
            except Exception as e:
                messages.error(request, f"Could not submit: {str(e)}")
        else:
            messages.error(request, "Please correct the errors in the form")
    
    return render(request, "staff_template/staff_feedback.html", context)


# =============================================================================
# PROFILE
# =============================================================================

@login_required
def staff_view_profile(request):
    """View and update profile"""
    if not check_faculty_permission(request.user):
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    user_form = AccountUserForm(request.POST or None, request.FILES or None, instance=request.user)
    profile_form = FacultyProfileEditForm(request.POST or None, instance=faculty)
    
    context = {
        'user_form': user_form,
        'profile_form': profile_form,
        'page_title': 'View/Update Profile'
    }
    
    if request.method == 'POST':
        if user_form.is_valid() and profile_form.is_valid():
            try:
                user = user_form.save(commit=False)
                password = user_form.cleaned_data.get('password')
                if password:
                    user.set_password(password)
                
                if 'profile_pic' in request.FILES:
                    fs = FileSystemStorage()
                    filename = fs.save(request.FILES['profile_pic'].name, request.FILES['profile_pic'])
                    user.profile_pic = fs.url(filename)
                
                user.save()
                profile_form.save()
                messages.success(request, "Profile Updated!")
                return redirect(reverse('staff_view_profile'))
            except Exception as e:
                messages.error(request, f"Error updating profile: {str(e)}")
        else:
            messages.error(request, "Invalid data provided")
    
    return render(request, "staff_template/staff_view_profile.html", context)


# =============================================================================
# NOTIFICATIONS
# =============================================================================

@csrf_exempt
@login_required
def staff_fcmtoken(request):
    """Update FCM token for push notifications"""
    token = request.POST.get('token')
    try:
        user = get_object_or_404(Account_User, id=request.user.id)
        user.fcm_token = token
        user.save()
        return HttpResponse("True")
    except Exception as e:
        return HttpResponse("False")


@login_required
def staff_view_notification(request):
    """View notifications"""
    if not check_faculty_permission(request.user):
        return redirect('/')
    
    notifications = Notification.objects.filter(recipient=request.user).order_by('-created_at')
    
    # Mark as read
    notifications.filter(is_read=False).update(is_read=True)
    
    context = {
        'notifications': notifications,
        'page_title': "View Notifications"
    }
    return render(request, "staff_template/staff_view_notification.html", context)


# =============================================================================
# PUBLICATIONS
# =============================================================================

@login_required
def staff_add_publication(request):
    """Add new publication"""
    if not check_faculty_permission(request.user):
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    form = PublicationForm(request.POST or None, request.FILES or None)
    publications = Publication.objects.filter(faculty=faculty).order_by('-created_at')
    
    context = {
        'form': form,
        'publications': publications,
        'page_title': 'Add Publication'
    }
    
    if request.method == 'POST':
        if form.is_valid():
            try:
                publication = form.save(commit=False)
                publication.faculty = faculty
                publication.save()
                messages.success(request, "Publication added successfully. Pending verification.")
                return redirect(reverse('staff_add_publication'))
            except Exception as e:
                messages.error(request, f"Could not add: {str(e)}")
        else:
            messages.error(request, "Please fill all required fields correctly")
    
    return render(request, "staff_template/staff_add_publication.html", context)


@login_required
def staff_view_publications(request):
    """View all publications"""
    if not check_faculty_permission(request.user):
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    publications = Publication.objects.filter(faculty=faculty).order_by('-created_at')
    
    context = {
        'publications': publications,
        'page_title': 'My Publications'
    }
    return render(request, "staff_template/staff_view_publications.html", context)


# =============================================================================
# VIEW STUDENTS
# =============================================================================

@login_required
def staff_view_students(request):
    """View students in assigned courses"""
    if not check_faculty_permission(request.user):
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    assignments = Course_Assignment.objects.filter(faculty=faculty, is_active=True)
    
    # Get unique batch labels
    batch_labels = assignments.values_list('batch_label', flat=True).distinct()
    
    # Get students in those batches
    students = Student_Profile.objects.filter(batch_label__in=batch_labels).select_related('user', 'advisor')
    
    context = {
        'students': students,
        'assignments': assignments,
        'page_title': 'View Students'
    }
    return render(request, "staff_template/staff_view_students.html", context)


# =============================================================================
# ATTENDANCE REPORTS
# =============================================================================

@login_required
def staff_view_attendance_report(request):
    """View attendance reports for assigned courses"""
    if not check_faculty_permission(request.user):
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    assignments = Course_Assignment.objects.filter(faculty=faculty, is_active=True)
    
    selected_assignment = None
    attendance_data = []
    
    if request.GET.get('assignment'):
        selected_assignment = get_object_or_404(Course_Assignment, id=request.GET.get('assignment'))
        
        # Get students and their attendance
        students = Student_Profile.objects.filter(batch_label=selected_assignment.batch_label)
        
        for student in students:
            total = Attendance.objects.filter(student=student, assignment=selected_assignment).count()
            present = Attendance.objects.filter(
                student=student, 
                assignment=selected_assignment,
                status='PRESENT'
            ).count()
            
            percentage = (present / total * 100) if total > 0 else 0
            
            attendance_data.append({
                'student': student,
                'total': total,
                'present': present,
                'absent': total - present,
                'percentage': round(percentage, 2)
            })
    
    context = {
        'assignments': assignments,
        'selected_assignment': selected_assignment,
        'attendance_data': attendance_data,
        'page_title': 'Attendance Report'
    }
    return render(request, "staff_template/staff_attendance_report.html", context)


# =============================================================================
# QUESTION PAPER MANAGEMENT
# =============================================================================

@login_required
def staff_view_qp_assignments(request):
    """View question paper assignments for the faculty"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied. Faculty privileges required.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    
    # Get all QP assignments for this faculty
    assignments = QuestionPaperAssignment.objects.filter(
        assigned_faculty=faculty
    ).select_related(
        'course', 'academic_year', 'semester', 'regulation'
    ).order_by('-created_at')
    
    # Statistics
    stats = {
        'total': assignments.count(),
        'pending': assignments.filter(status__in=['ASSIGNED', 'IN_PROGRESS']).count(),
        'submitted': assignments.filter(status='SUBMITTED').count(),
        'approved': assignments.filter(status='APPROVED').count(),
        'rejected': assignments.filter(status='REJECTED').count(),
    }
    
    context = {
        'assignments': assignments,
        'stats': stats,
        'page_title': 'My Question Paper Assignments'
    }
    return render(request, "staff_template/staff_qp_assignments.html", context)


@login_required
def staff_submit_question_paper(request, qp_id):
    """Submit question paper for a specific assignment"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied. Faculty privileges required.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    qp_assignment = get_object_or_404(QuestionPaperAssignment, id=qp_id, assigned_faculty=faculty)
    
    # Check if already approved
    if qp_assignment.status == 'APPROVED':
        messages.info(request, "This question paper has already been approved.")
        return redirect('staff_view_qp_assignments')
    
    form = QuestionPaperSubmissionForm(request.POST or None, request.FILES or None, instance=qp_assignment)
    
    context = {
        'form': form,
        'qp_assignment': qp_assignment,
        'page_title': f'Submit QP - {qp_assignment.course.course_code}'
    }
    
    if request.method == 'POST':
        if form.is_valid():
            try:
                qp = form.save(commit=False)
                qp.status = 'SUBMITTED'
                qp.submitted_at = datetime.now()
                qp.save()
                
                # Create notification for HOD
                hod_users = Account_User.objects.filter(role='HOD', is_active=True)
                for hod in hod_users:
                    Notification.objects.create(
                        recipient=hod,
                        title='Question Paper Submitted',
                        message=f'{faculty.user.full_name} has submitted the {qp.get_exam_type_display()} question paper for {qp.course.course_code} - {qp.course.title}',
                        notification_type='INFO',
                        link=reverse('review_question_paper', args=[qp.id])
                    )
                
                messages.success(request, "Question paper submitted successfully! Waiting for review.")
                return redirect('staff_view_qp_assignments')
            except Exception as e:
                messages.error(request, f"Error submitting: {str(e)}")
        else:
            messages.error(request, "Please correct the errors in the form")
    
    return render(request, "staff_template/staff_submit_qp.html", context)


@login_required
def staff_view_qp_details(request, qp_id):
    """View details of a specific question paper assignment"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied. Faculty privileges required.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    qp_assignment = get_object_or_404(QuestionPaperAssignment, id=qp_id, assigned_faculty=faculty)
    
    context = {
        'qp_assignment': qp_assignment,
        'page_title': f'QP Details - {qp_assignment.course.course_code}'
    }
    return render(request, "staff_template/staff_qp_details.html", context)


# =============================================================================
# TIMETABLE VIEW
# =============================================================================

@login_required
def staff_view_timetable(request):
    """View faculty's teaching schedule across all batches"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied. Faculty privileges required.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    
    # Get all timetable entries for this faculty (as main faculty or lab assistant).
    direct_entries = TimetableEntry.objects.filter(
        Q(faculty=faculty) | Q(lab_assistant=faculty)
    ).select_related(
        'timetable__academic_year', 'timetable__semester',
        'course', 'time_slot', 'lab_room'
    ).order_by('timetable', 'day', 'time_slot__slot_number')

    # Include placeholder elective slots where faculty is mapped via offering assignment
    # (PEC/OEC placeholders may store a representative faculty in TimetableEntry).
    direct_entry_ids = set(direct_entries.values_list('id', flat=True))

    offering_rows = ElectiveOfferingFacultyAssignment.objects.filter(
        faculty=faculty,
        is_active=True,
        offering__is_active=True,
        offering__regulation_course_plan__course__is_placeholder=True,
    ).values_list(
        'offering__semester_id',
        'offering__regulation_course_plan__course_id',
        'offering__regulation_course_plan__branch',
        'offering__regulation_course_plan__program_type',
        'offering__semester__semester_number',
        'offering__actual_course__course_code',
        'offering__actual_course__title',
        'batch_number',
    ).distinct()

    mapped_semester_ids = set()
    mapped_placeholder_ids = set()
    mapped_exact_keys = set()
    mapped_loose_keys = set()
    mapped_codes_exact = {}
    mapped_codes_loose = {}
    mapped_titles_exact = {}
    mapped_titles_loose = {}
    mapped_batch_numbers_exact = {}
    mapped_batch_names_exact = {}
    batch_name_order_cache = {}

    for semester_id, placeholder_id, branch, program_type, semester_number, actual_course_code, actual_course_title, batch_number in offering_rows:
        mapped_semester_ids.add(semester_id)
        mapped_placeholder_ids.add(placeholder_id)
        mapped_loose_keys.add((semester_id, placeholder_id))
        year_of_study = (int(semester_number) + 1) // 2 if semester_number else None
        exact_key = (
            semester_id,
            placeholder_id,
            (branch or '').strip().upper(),
            (program_type or '').strip().upper(),
            year_of_study,
        )
        loose_key = (semester_id, placeholder_id)
        mapped_exact_keys.add(exact_key)

        if exact_key not in mapped_codes_exact:
            mapped_codes_exact[exact_key] = set()
        if exact_key not in mapped_titles_exact:
            mapped_titles_exact[exact_key] = {}
        if exact_key not in mapped_batch_numbers_exact:
            mapped_batch_numbers_exact[exact_key] = set()
        if loose_key not in mapped_codes_loose:
            mapped_codes_loose[loose_key] = set()
        if loose_key not in mapped_titles_loose:
            mapped_titles_loose[loose_key] = {}

        if actual_course_code:
            code = actual_course_code.strip().upper()
            mapped_codes_exact[exact_key].add(code)
            mapped_codes_loose[loose_key].add(code)
            if actual_course_title:
                mapped_titles_exact[exact_key][code] = actual_course_title
                mapped_titles_loose[loose_key][code] = actual_course_title

        if batch_number:
            mapped_batch_numbers_exact[exact_key].add(int(batch_number))

    # Resolve elective batch_number assignments (1-based) to concrete batch names
    # using timetable batch ordering for the same semester/program/year.
    for exact_key, batch_numbers in mapped_batch_numbers_exact.items():
        sem_id, _placeholder_id, branch_code, program_level, year_of_study = exact_key
        cache_key = (sem_id, branch_code, program_level, year_of_study)

        if cache_key not in batch_name_order_cache:
            ordered_batch_names = list(
                Timetable.objects.filter(
                    is_active=True,
                    semester_id=sem_id,
                    year=year_of_study,
                    program_batch__program__code__iexact=branch_code,
                    program_batch__program__level__iexact=program_level,
                ).order_by('program_batch__batch_name').values_list('program_batch__batch_name', flat=True).distinct()
            )
            batch_name_order_cache[cache_key] = ordered_batch_names

        ordered_names = batch_name_order_cache.get(cache_key, [])
        resolved_names = set()
        for bn in batch_numbers:
            idx = bn - 1
            if 0 <= idx < len(ordered_names):
                resolved_names.add(ordered_names[idx])
        mapped_batch_names_exact[exact_key] = resolved_names

    indirect_entry_ids = set()
    if mapped_semester_ids and mapped_placeholder_ids:
        candidate_entries = TimetableEntry.objects.filter(
            course__is_placeholder=True,
            timetable__semester_id__in=mapped_semester_ids,
            course_id__in=mapped_placeholder_ids,
        ).exclude(
            id__in=direct_entry_ids
        ).select_related(
            'timetable__program_batch__program'
        )

        for entry in candidate_entries:
            key_loose = (entry.timetable.semester_id, entry.course_id)
            if key_loose not in mapped_loose_keys:
                continue

            program = entry.timetable.program_batch.program if entry.timetable.program_batch_id and entry.timetable.program_batch and entry.timetable.program_batch.program else None
            if program:
                key_exact = (
                    entry.timetable.semester_id,
                    entry.course_id,
                    (program.code or '').strip().upper(),
                    (program.level or '').strip().upper(),
                    entry.timetable.year,
                )
                if key_exact in mapped_exact_keys:
                    allowed_batch_names = mapped_batch_names_exact.get(key_exact, set())
                    current_batch = entry.timetable.batch_display
                    if allowed_batch_names and current_batch not in allowed_batch_names:
                        continue
                    indirect_entry_ids.add(entry.id)
            else:
                # Skip legacy timetables without program_batch linkage to avoid
                # overcounting elective slots from unrelated batches.
                continue

    all_entry_ids = direct_entry_ids | indirect_entry_ids
    entries = list(TimetableEntry.objects.filter(id__in=all_entry_ids).select_related(
        'timetable__academic_year', 'timetable__semester',
        'course', 'time_slot', 'lab_room', 'timetable__program_batch__program'
    ).order_by('timetable', 'day', 'time_slot__slot_number'))
    
    # Precompute IOC course codes per timetable so timetable display can show "IOC".
    timetable_ioc_codes = {}
    for tt in {entry.timetable for entry in entries}:
        quick_ioc_qs = Course_Assignment.objects.filter(
            semester=tt.semester,
            academic_year=tt.academic_year,
            special_note__icontains='Quick IOC',
            is_active=True,
        )

        if tt.program_batch_id:
            batch = tt.program_batch
            quick_ioc_qs = quick_ioc_qs.filter(batch=batch)
        else:
            batch = ProgramBatch.objects.filter(
                academic_year=tt.academic_year,
                year_of_study=tt.year,
                batch_name=tt.batch,
                is_active=True,
            ).select_related('program').first()
            if batch:
                quick_ioc_qs = quick_ioc_qs.filter(batch=batch)
            else:
                quick_ioc_qs = quick_ioc_qs.filter(batch__batch_name=tt.batch, batch__year_of_study=tt.year)

        ioc_codes = set(quick_ioc_qs.values_list('course__course_code', flat=True))

        if batch and batch.program:
            ioc_codes.update(
                RegulationCoursePlan.objects.filter(
                    semester=tt.semester.semester_number,
                    branch=batch.program.code,
                    program_type=batch.program.level,
                ).filter(
                    Q(category__code__in=['IOC', 'EEC']) |
                    Q(course__placeholder_type__in=['IOC', 'EEC'])
                ).values_list('course__course_code', flat=True)
            )

        timetable_ioc_codes[tt.id] = ioc_codes

    for entry in entries:
        entry.hide_batch_details = False
        entry.include_in_detailed = True
        entry.include_in_consolidated = True
        entry.display_title = entry.course.title if entry.course else None

        if entry.course:
            if entry.course.is_placeholder:
                mapped_codes = set()
                mapped_titles = {}
                loose_key = (entry.timetable.semester_id, entry.course_id)
                program = entry.timetable.program_batch.program if entry.timetable.program_batch_id and entry.timetable.program_batch and entry.timetable.program_batch.program else None

                if program:
                    exact_key = (
                        entry.timetable.semester_id,
                        entry.course_id,
                        (program.code or '').strip().upper(),
                        (program.level or '').strip().upper(),
                        entry.timetable.year,
                    )
                    mapped_codes = set(mapped_codes_exact.get(exact_key, set()))
                    mapped_titles = dict(mapped_titles_exact.get(exact_key, {}))

                    allowed_batch_names = mapped_batch_names_exact.get(exact_key, set())
                    current_batch = entry.timetable.batch_display
                    if allowed_batch_names and current_batch not in allowed_batch_names:
                        entry.display_code = None
                        entry.include_in_detailed = False
                        entry.include_in_consolidated = False
                        continue

                if not mapped_codes:
                    mapped_codes = set(mapped_codes_loose.get(loose_key, set()))
                    mapped_titles = dict(mapped_titles_loose.get(loose_key, {}))

                if mapped_codes:
                    slot_codes = set()
                    if entry.special_note:
                        raw_note = str(entry.special_note).replace(',', '/').replace('|', '/')
                        for token in raw_note.split('/'):
                            token = token.strip().upper()
                            if token:
                                slot_codes.add(token)

                    # For PEC/OEC placeholder rows, show only slots that actually
                    # contain this faculty's assigned actual elective code.
                    display_codes = set(mapped_codes)
                    if slot_codes:
                        display_codes = mapped_codes.intersection(slot_codes)
                        if not display_codes:
                            entry.display_code = None
                            entry.include_in_detailed = False
                            entry.include_in_consolidated = False
                            continue

                    # For PEC/OEC slots, faculty should see only their assigned actual course(s).
                    entry.display_code = '/'.join(sorted(display_codes))
                    display_titles = [mapped_titles.get(code) for code in sorted(display_codes) if mapped_titles.get(code)]
                    if display_titles:
                        entry.display_title = ' / '.join(display_titles)
                    entry.hide_batch_details = True
                    entry.include_in_detailed = True
                else:
                    codes = timetable_ioc_codes.get(entry.timetable_id, set())
                    is_ioc = (
                        entry.course.course_code in codes or
                        (entry.course.is_placeholder and entry.course.placeholder_type in ['IOC', 'EEC'])
                    )
                    entry.display_code = 'IOC' if is_ioc else entry.course.course_code
            else:
                codes = timetable_ioc_codes.get(entry.timetable_id, set())
                is_ioc = (
                    entry.course.course_code in codes or
                    (entry.course.is_placeholder and entry.course.placeholder_type in ['IOC', 'EEC'])
                )
                entry.display_code = 'IOC' if is_ioc else entry.course.course_code
        else:
            entry.display_code = None

        if entry.hide_batch_details:
            entry.display_batch_label = f"Y{entry.timetable.year}"
        else:
            entry.display_batch_label = f"Y{entry.timetable.year}-{entry.timetable.batch}"

    # Group entries by timetable (batch)
    timetable_entries = {}
    for entry in entries:
        if not entry.include_in_detailed:
            continue

        tt_key = entry.timetable.id
        if tt_key not in timetable_entries:
            timetable_entries[tt_key] = {
                'timetable': entry.timetable,
                'entries': [],
                'all_hide_batch_details': True,
            }
        timetable_entries[tt_key]['entries'].append(entry)
        if not getattr(entry, 'hide_batch_details', False):
            timetable_entries[tt_key]['all_hide_batch_details'] = False

    for tt_data in timetable_entries.values():
        tt = tt_data['timetable']
        if tt_data.get('all_hide_batch_details'):
            tt_data['display_batch_header'] = f"Year {tt.year}"
        else:
            tt_data['display_batch_header'] = f"Year {tt.year} - Batch {tt.batch}"
    
    # Get time slots and days for consolidated view
    time_slots = TimeSlot.objects.all().order_by('slot_number')
    days = TimetableEntry.DAY_CHOICES
    
    # Create consolidated schedule (all batches combined)
    entries_by_slot = defaultdict(list)
    for entry in entries:
        if not getattr(entry, 'include_in_consolidated', True):
            continue
        entries_by_slot[(entry.day, entry.time_slot.slot_number)].append(entry)

    consolidated_schedule = {}
    for day_code, day_name in days:
        consolidated_schedule[day_code] = {}
        for slot in time_slots:
            raw_entries = entries_by_slot.get((day_code, slot.slot_number), [])
            deduped = []
            seen_pec = set()
            for entry in raw_entries:
                if entry.hide_batch_details and entry.display_code:
                    dedupe_key = (entry.display_code, entry.timetable.year)
                    if dedupe_key in seen_pec:
                        continue
                    seen_pec.add(dedupe_key)
                deduped.append(entry)
            consolidated_schedule[day_code][slot.slot_number] = deduped
    
    context = {
        'faculty': faculty,
        'timetable_entries': timetable_entries,
        'consolidated_schedule': consolidated_schedule,
        'time_slots': time_slots,
        'days': days,
        'page_title': 'My Teaching Schedule'
    }
    return render(request, "staff_template/staff_timetable.html", context)


# =============================================================================
# STRUCTURED QUESTION PAPER (R2023 Format - Multi-field)
# =============================================================================

@login_required
def staff_list_structured_qps(request):
    """List all structured question papers and assignments for faculty"""
    from datetime import date
    
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied. Faculty privileges required.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    
    # Get faculty's QPs
    qps = StructuredQuestionPaper.objects.filter(faculty=faculty).select_related(
        'course', 'academic_year', 'semester', 'regulation'
    ).order_by('-created_at')
    
    # Get pending assignments (not yet completed, including rejected for revision)
    pending_assignments = QuestionPaperAssignment.objects.filter(
        assigned_faculty=faculty,
        status__in=['ASSIGNED', 'IN_PROGRESS', 'REJECTED', 'REVISION_REQUIRED']
    ).select_related('course', 'academic_year', 'semester', 'regulation'
    ).prefetch_related('structured_qp').order_by('deadline')
    
    # Count stats for dashboard
    draft_count = qps.filter(status='DRAFT').count()
    submitted_count = qps.filter(status='SUBMITTED').count()
    approved_count = qps.filter(status='APPROVED').count()
    uploaded_count = qps.filter(is_uploaded=True).count()
    
    context = {
        'qps': qps,
        'pending_assignments': pending_assignments,
        'draft_count': draft_count,
        'submitted_count': submitted_count,
        'approved_count': approved_count,
        'uploaded_count': uploaded_count,
        'page_title': 'Question Paper Dashboard'
    }
    return render(request, "staff_template/list_structured_qps.html", context)


@login_required
def staff_upload_qp(request):
    """Upload a question paper document directly"""
    from main_app.forms import UploadQuestionPaperForm
    
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied. Faculty privileges required.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    
    if request.method == 'POST':
        form = UploadQuestionPaperForm(request.POST, request.FILES, faculty=faculty)
        if form.is_valid():
            try:
                qp = StructuredQuestionPaper(
                    faculty=faculty,
                    course=form.cleaned_data['course'],
                    academic_year=form.cleaned_data['academic_year'],
                    semester=form.cleaned_data['semester'],
                    regulation=form.cleaned_data['regulation'],
                    exam_month_year=form.cleaned_data['exam_month_year'],
                    uploaded_document=form.cleaned_data['uploaded_document'],
                    is_uploaded=True,
                    status='DRAFT'
                )
                qp.save()
                
                messages.success(request, f"Question paper uploaded successfully for {qp.course.course_code}!")
                return redirect('staff_list_structured_qps')
            except Exception as e:
                messages.error(request, f"Error uploading question paper: {str(e)}")
        else:
            messages.error(request, "Please correct the errors in the form.")
    else:
        form = UploadQuestionPaperForm(faculty=faculty)
    
    context = {
        'form': form,
        'page_title': 'Upload Question Paper'
    }
    return render(request, "staff_template/upload_qp.html", context)


@login_required
def staff_create_structured_qp(request, assignment_id=None):
    """Create structured question paper with formsets"""
    from main_app.forms import StructuredQuestionPaperForm, PartAFormSet, PartBFormSet, PartCFormSet
    
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied. Faculty privileges required.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    assignment = None
    
    if assignment_id:
        assignment = get_object_or_404(QuestionPaperAssignment, id=assignment_id, assigned_faculty=faculty)
        if hasattr(assignment, 'structured_qp'):
            messages.info(request, "Structured question paper already exists.")
            return redirect('staff_edit_structured_qp', qp_id=assignment.structured_qp.id)
    
    if request.method == 'POST':
        qp_form = StructuredQuestionPaperForm(request.POST)
        
        if qp_form.is_valid():
            qp = qp_form.save(commit=False)
            qp.faculty = faculty
            qp.qp_assignment = assignment
            
            # Validate formsets BEFORE saving QP (use a temporary unsaved instance)
            part_a_formset = PartAFormSet(request.POST, request.FILES, instance=qp, prefix='part_a',
                                         min_required=10, part_label='Part A')
            part_b_formset = PartBFormSet(request.POST, request.FILES, instance=qp, prefix='part_b',
                                         min_required=10, part_label='Part B')
            part_c_formset = PartCFormSet(request.POST, request.FILES, instance=qp, prefix='part_c',
                                         min_required=1, part_label='Part C')
            
            a_valid = part_a_formset.is_valid()
            b_valid = part_b_formset.is_valid()
            c_valid = part_c_formset.is_valid()
            
            if a_valid and b_valid and c_valid:
                # All good — now persist the QP
                qp.save()
                
                # Re-bind formsets to the saved instance so FK is set
                part_a_formset = PartAFormSet(request.POST, request.FILES, instance=qp, prefix='part_a',
                                             min_required=10, part_label='Part A')
                part_b_formset = PartBFormSet(request.POST, request.FILES, instance=qp, prefix='part_b',
                                             min_required=10, part_label='Part B')
                part_c_formset = PartCFormSet(request.POST, request.FILES, instance=qp, prefix='part_c',
                                             min_required=1, part_label='Part C')
                part_a_formset.is_valid()
                part_b_formset.is_valid()
                part_c_formset.is_valid()
                
                # Save Part A questions
                for i, form in enumerate(part_a_formset):
                    if form.cleaned_data and form.cleaned_data.get('question_text'):
                        question = form.save(commit=False)
                        question.question_paper = qp
                        question.part = 'A'
                        question.question_number = i + 1
                        question.marks = 2
                        question.save()
                
                # Save Part B questions
                for i, form in enumerate(part_b_formset):
                    if form.cleaned_data and form.cleaned_data.get('question_text'):
                        question = form.save(commit=False)
                        question.question_paper = qp
                        question.part = 'B'
                        question.is_or_option = True
                        question.or_pair_number = 11 + (i // 2)
                        question.option_label = '(a)' if i % 2 == 0 else '(b)'
                        question.question_number = 11 + (i // 2)
                        question.marks = 13
                        question.save()
                
                # Save Part C question
                for form in part_c_formset:
                    if form.cleaned_data and form.cleaned_data.get('question_text'):
                        question = form.save(commit=False)
                        question.question_paper = qp
                        question.part = 'C'
                        question.question_number = 16
                        question.marks = 15
                        question.save()
                
                # Validate BL% distribution before allowing preview
                validation = qp.validate_distribution()
                if validation['errors']:
                    for err in validation['errors']:
                        messages.error(request, err)
                    for sug in validation.get('suggestions', []):
                        messages.warning(request, sug)
                    return redirect('staff_edit_structured_qp', qp_id=qp.id)
                
                messages.success(request, "Question paper created successfully!")
                return redirect('staff_preview_structured_qp', qp_id=qp.id)
            else:
                # Formset validation failed — collect all errors
                all_errors = []
                for label, fs in [('Part A', part_a_formset), ('Part B', part_b_formset), ('Part C', part_c_formset)]:
                    # Non-form errors (min-count check)
                    for err in fs.non_form_errors():
                        all_errors.append(str(err))
                    # Per-form errors
                    for idx, form in enumerate(fs):
                        for field, errs in form.errors.items():
                            for e in errs:
                                q_num = idx + 1 if label != 'Part B' else idx + 11
                                all_errors.append(f'{label} Q{q_num} — {field}: {e}')
                
                if all_errors:
                    for err_msg in all_errors[:10]:   # Show first 10 errors
                        messages.error(request, err_msg)
                    if len(all_errors) > 10:
                        messages.error(request, f'… and {len(all_errors) - 10} more error(s).')
                else:
                    messages.error(request, "Please check the questions for errors.")
        else:
            part_a_formset = PartAFormSet(request.POST, request.FILES, prefix='part_a',
                                         min_required=10, part_label='Part A')
            part_b_formset = PartBFormSet(request.POST, request.FILES, prefix='part_b',
                                         min_required=10, part_label='Part B')
            part_c_formset = PartCFormSet(request.POST, request.FILES, prefix='part_c',
                                         min_required=1, part_label='Part C')
            messages.error(request, "Please fix the basic details errors above.")
    else:
        initial_data = {}
        if assignment:
            initial_data = {
                'course': assignment.course,
                'academic_year': assignment.academic_year,
                'semester': assignment.semester,
                'regulation': assignment.regulation,
            }
        
        qp_form = StructuredQuestionPaperForm(initial=initial_data)
        
        # Auto-populate formsets with empty forms
        part_a_initial = [{'question_number': i+1, 'marks': 2, 'part': 'A'} for i in range(10)]
        part_b_initial = [{'marks': 13, 'part': 'B'} for i in range(10)]
        part_c_initial = [{'question_number': 16, 'marks': 15, 'part': 'C'}]
        
        part_a_formset = PartAFormSet(prefix='part_a', queryset=QPQuestion.objects.none(),
                                     min_required=10, part_label='Part A')
        part_b_formset = PartBFormSet(prefix='part_b', queryset=QPQuestion.objects.none(),
                                     min_required=10, part_label='Part B')
        part_c_formset = PartCFormSet(prefix='part_c', queryset=QPQuestion.objects.none(),
                                     min_required=1, part_label='Part C')
    
    context = {
        'qp_form': qp_form,
        'part_a_formset': part_a_formset,
        'part_b_formset': part_b_formset,
        'part_c_formset': part_c_formset,
        'assignment': assignment,
        'page_title': 'Create Structured Question Paper'
    }
    return render(request, "staff_template/create_structured_qp.html", context)


@login_required
def staff_create_qp_from_upload(request, assignment_id=None):
    """
    Create structured QP from PDF/DOCX file upload with auto-extraction.
    
    Flow:
    1. Faculty uploads PDF/DOCX file
    2. System auto-extracts questions from file
    3. Display extracted data for review/manual mapping of CO and Bloom's levels
    4. Save extracted questions to StructuredQuestionPaper
    5. Run validation
    6. Faculty can then submit for review
    """
    from main_app.forms import UploadedQPForm
    from main_app.qp_extraction import extract_qp_from_file, create_qp_questions_from_extraction
    
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied. Faculty privileges required.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    assignment = None
    
    if assignment_id:
        assignment = get_object_or_404(QuestionPaperAssignment, id=assignment_id, assigned_faculty=faculty)
        if hasattr(assignment, 'structured_qp'):
            messages.info(request, "Structured question paper already exists for this assignment.")
            return redirect('staff_preview_structured_qp', qp_id=assignment.structured_qp.id)
    
    # Get step from GET params (preferred) or POST params (form submission), default to 1
    step = request.GET.get('step') or request.POST.get('step') or '1'
    step = str(step)  # Ensure it's a string for comparison
    
    if request.method == 'POST':
        if step == '1':  # Upload & Extract
            form = UploadedQPForm(request.POST, request.FILES)
            
            if form.is_valid():
                uploaded_file = form.cleaned_data['qp_file']

                # Save QP (not yet with questions)
                qp = form.save(commit=False)
                qp.faculty = faculty
                qp.qp_assignment = assignment
                qp.is_uploaded = True
                qp.uploaded_document = uploaded_file
                qp.status = 'DRAFT'
                qp.save()
                
                # Extract questions from uploaded file
                extraction_result = extract_qp_from_file(uploaded_file, uploaded_file.name)

                # Auto-fill CO descriptions from extracted QP content if user left them blank.
                extracted_co = extraction_result.get('co_descriptions', {}) or {}
                co_fields = {
                    'co1_description': extracted_co.get('CO1', ''),
                    'co2_description': extracted_co.get('CO2', ''),
                    'co3_description': extracted_co.get('CO3', ''),
                    'co4_description': extracted_co.get('CO4', ''),
                    'co5_description': extracted_co.get('CO5', ''),
                }
                fields_to_update = []
                for field_name, extracted_value in co_fields.items():
                    if extracted_value and not getattr(qp, field_name, '').strip():
                        setattr(qp, field_name, extracted_value)
                        fields_to_update.append(field_name)

                if fields_to_update:
                    qp.save(update_fields=fields_to_update)
                
                # Store extraction result in session for review
                request.session[f'qp_{qp.id}_extraction'] = {
                    'part_a': extraction_result.get('part_a', []),
                    'part_b': extraction_result.get('part_b', []),
                    'part_c': extraction_result.get('part_c', []),
                    'co_descriptions': extraction_result.get('co_descriptions', {}),
                    'errors': extraction_result.get('errors', []),
                    'total_marks': extraction_result.get('total_marks', 0)
                }
                request.session[f'qp_{qp.id}_file_info'] = {
                    'filename': uploaded_file.name,
                    'size': uploaded_file.size
                }
                
                if extraction_result.get('errors'):
                    for error in extraction_result['errors']:
                        messages.warning(request, f"Extraction notice: {error}")
                
                messages.success(request, f"File uploaded and questions extracted! Total marks: {extraction_result.get('total_marks', 0)}")
                return redirect(f"{reverse('staff_create_qp_from_upload')}?step=2&qp_id={qp.id}&assignment_id={assignment_id or ''}")
            else:
                for field, errors in form.errors.items():
                    for error in errors:
                        messages.error(request, f"{field}: {error}")
                
                # Render form again with errors
                context = {
                    'form': form,
                    'step': 1,
                    'assignment': assignment,
                    'page_title': 'Create QP from File Upload'
                }
                return render(request, "staff_template/create_qp_from_upload.html", context)
        
        elif step == '2':  # Review & Map CO/Bloom's
            qp_id = request.POST.get('qp_id')
            qp = get_object_or_404(StructuredQuestionPaper, id=qp_id, faculty=faculty)
            
            extraction_key = f'qp_{qp.id}_extraction'
            if extraction_key not in request.session:
                messages.error(request, "Extraction data not found. Please upload again.")
                return redirect('staff_create_qp_from_upload', assignment_id=assignment_id)
            
            extraction = request.session[extraction_key]
            
            # Extract CO and Bloom's mappings from POST or use extracted values
            co_mapping = {}
            bloom_mapping = {}
            
            # Parse Part A mappings - use POST values if provided, otherwise use extracted values
            for idx, q in enumerate(extraction.get('part_a', [])):
                co_key = f'part_a_{idx}_co'
                bloom_key = f'part_a_{idx}_bloom'
                # Use POST value if submitted (user changed it), otherwise use extracted value
                co_mapping[f'a_{idx}'] = request.POST.get(co_key) or q.get('co', 'CO1')
                bloom_mapping[f'a_{idx}'] = request.POST.get(bloom_key) or q.get('bl', 'L1')
            
            # Parse Part B mappings
            for idx, q in enumerate(extraction.get('part_b', [])):
                co_key = f'part_b_{idx}_co'
                bloom_key = f'part_b_{idx}_bloom'
                co_mapping[f'b_{idx}'] = request.POST.get(co_key) or q.get('co', 'CO1')
                bloom_mapping[f'b_{idx}'] = request.POST.get(bloom_key) or q.get('bl', 'L1')
            
            # Parse Part C mappings
            for idx, q in enumerate(extraction.get('part_c', [])):
                co_key = f'part_c_{idx}_co'
                bloom_key = f'part_c_{idx}_bloom'
                co_mapping[f'c_{idx}'] = request.POST.get(co_key) or q.get('co', 'CO1')
                bloom_mapping[f'c_{idx}'] = request.POST.get(bloom_key) or q.get('bl', 'L1')
            
            # Create question objects from extraction
            questions_created, creation_errors = create_qp_questions_from_extraction(
                qp, extraction, bloom_mapping, co_mapping
            )
            
            if creation_errors:
                for error in creation_errors:
                    messages.warning(request, error)
            
            if questions_created > 0:
                messages.success(request, f"Created {questions_created} questions from uploaded file!")
                
                # Validate distribution
                validation = qp.validate_distribution()
                if validation['errors']:
                    for err in validation['errors']:
                        messages.error(request, err)
                    for sug in validation.get('suggestions', []):
                        messages.warning(request, sug)
                else:
                    messages.success(request, "Question distribution looks good!")
                
                # Cleanup session
                if extraction_key in request.session:
                    del request.session[extraction_key]
                
                return redirect('staff_preview_structured_qp', qp_id=qp.id)
            else:
                messages.error(request, "Failed to create questions from extracted data.")
                return redirect('staff_preview_structured_qp', qp_id=qp.id)
    
    else:  # GET request
        if step == '2':  # Review step
            qp_id = request.GET.get('qp_id')
            qp = get_object_or_404(StructuredQuestionPaper, id=qp_id, faculty=faculty)
            
            extraction_key = f'qp_{qp.id}_extraction'
            if extraction_key not in request.session:
                messages.error(request, "Extraction data not found. Please upload again.")
                return redirect('staff_create_qp_from_upload', assignment_id=assignment_id)
            
            extraction = request.session[extraction_key]
            file_info = request.session.get(f'qp_{qp.id}_file_info', {})
            
            context = {
                'qp': qp,
                'step': 2,
                'extraction': extraction,
                'file_info': file_info,
                'co_choices': ['CO1', 'CO2', 'CO3', 'CO4', 'CO5'],
                'bloom_choices': ['L1', 'L2', 'L3', 'L4', 'L5', 'L6'],
                'assignment': assignment,
                'page_title': f'Review Extracted QP - {qp.course.course_code}'
            }
            return render(request, "staff_template/create_qp_from_upload_step2.html", context)
        
        else:  # Step 1: Upload form
            initial_data = {}
            if assignment:
                initial_data = {
                    'course': assignment.course,
                    'academic_year': assignment.academic_year,
                    'semester': assignment.semester,
                    'regulation': assignment.regulation,
                }
            
            form = UploadedQPForm(initial=initial_data)
            
            context = {
                'form': form,
                'step': 1,
                'assignment': assignment,
                'page_title': 'Create QP from File Upload'
            }
            return render(request, "staff_template/create_qp_from_upload.html", context)


@login_required
def staff_edit_structured_qp(request, qp_id):
    """Edit existing structured question paper"""
    from main_app.forms import StructuredQuestionPaperForm, PartAFormSet, PartBFormSet, PartCFormSet
    
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    qp = get_object_or_404(StructuredQuestionPaper, id=qp_id, faculty=faculty)
    
    if qp.status not in ['DRAFT', 'REJECTED']:
        messages.warning(request, "Cannot edit submitted question paper.")
        return redirect('staff_preview_structured_qp', qp_id=qp.id)
    
    if request.method == 'POST':
        qp_form = StructuredQuestionPaperForm(request.POST, instance=qp)
        part_a_formset = PartAFormSet(request.POST, request.FILES, instance=qp, prefix='part_a',
                                     queryset=qp.questions.filter(part='A').order_by('question_number'),
                                     min_required=10, part_label='Part A')
        part_b_formset = PartBFormSet(request.POST, request.FILES, instance=qp, prefix='part_b',
                                     queryset=qp.questions.filter(part='B').order_by('or_pair_number', 'option_label'),
                                     min_required=10, part_label='Part B')
        part_c_formset = PartCFormSet(request.POST, request.FILES, instance=qp, prefix='part_c',
                                     queryset=qp.questions.filter(part='C'),
                                     min_required=1, part_label='Part C')
        
        a_valid = part_a_formset.is_valid()
        b_valid = part_b_formset.is_valid()
        c_valid = part_c_formset.is_valid()
        
        if qp_form.is_valid() and a_valid and b_valid and c_valid:
            qp_form.save()
            
            # Update questions
            for i, form in enumerate(part_a_formset):
                if form.cleaned_data and form.cleaned_data.get('question_text'):
                    question = form.save(commit=False)
                    question.part = 'A'
                    question.question_number = i + 1
                    question.marks = 2
                    question.save()
            
            for i, form in enumerate(part_b_formset):
                if form.cleaned_data and form.cleaned_data.get('question_text'):
                    question = form.save(commit=False)
                    question.part = 'B'
                    question.is_or_option = True
                    question.or_pair_number = 11 + (i // 2)
                    question.option_label = '(a)' if i % 2 == 0 else '(b)'
                    question.question_number = 11 + (i // 2)
                    question.marks = 13
                    question.save()
            
            for form in part_c_formset:
                if form.cleaned_data and form.cleaned_data.get('question_text'):
                    question = form.save(commit=False)
                    question.part = 'C'
                    question.question_number = 16
                    question.marks = 15
                    question.save()
            
            # Validate BL% distribution before allowing preview
            validation = qp.validate_distribution()
            if validation['errors']:
                for err in validation['errors']:
                    messages.error(request, err)
                for sug in validation.get('suggestions', []):
                    messages.warning(request, sug)
                return redirect('staff_edit_structured_qp', qp_id=qp.id)
            
            messages.success(request, "Question paper updated successfully!")
            return redirect('staff_preview_structured_qp', qp_id=qp.id)
        else:
            # Collect detailed errors
            all_errors = []
            if qp_form.errors:
                for field, errs in qp_form.errors.items():
                    for e in errs:
                        all_errors.append(f'Basic Details — {field}: {e}')
            for label, fs in [('Part A', part_a_formset), ('Part B', part_b_formset), ('Part C', part_c_formset)]:
                for err in fs.non_form_errors():
                    all_errors.append(str(err))
                for idx, form in enumerate(fs):
                    for field, errs in form.errors.items():
                        for e in errs:
                            q_num = idx + 1 if label != 'Part B' else idx + 11
                            all_errors.append(f'{label} Q{q_num} — {field}: {e}')
            for err_msg in all_errors[:10]:
                messages.error(request, err_msg)
            if len(all_errors) > 10:
                messages.error(request, f'… and {len(all_errors) - 10} more error(s).')
    else:
        qp_form = StructuredQuestionPaperForm(instance=qp)
        part_a_formset = PartAFormSet(instance=qp, prefix='part_a',
                                     queryset=qp.questions.filter(part='A').order_by('question_number'),
                                     min_required=10, part_label='Part A')
        part_b_formset = PartBFormSet(instance=qp, prefix='part_b',
                                     queryset=qp.questions.filter(part='B').order_by('or_pair_number', 'option_label'),
                                     min_required=10, part_label='Part B')
        part_c_formset = PartCFormSet(instance=qp, prefix='part_c',
                                     queryset=qp.questions.filter(part='C'),
                                     min_required=1, part_label='Part C')
    
    context = {
        'qp_form': qp_form,
        'part_a_formset': part_a_formset,
        'part_b_formset': part_b_formset,
        'part_c_formset': part_c_formset,
        'qp': qp,
        'page_title': 'Edit Structured Question Paper'
    }
    return render(request, "staff_template/create_structured_qp.html", context)


@login_required
def staff_preview_structured_qp(request, qp_id):
    """Preview structured question paper"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    qp = get_object_or_404(StructuredQuestionPaper, id=qp_id, faculty=faculty)
    
    # Get questions by part
    part_a_questions = qp.questions.filter(part='A').order_by('question_number')
    part_b_questions = qp.questions.filter(part='B').order_by('or_pair_number', 'option_label')
    part_c_questions = qp.questions.filter(part='C')
    
    # Group Part B into OR pairs
    part_b_pairs = []
    for i in range(11, 16):
        pair_questions = part_b_questions.filter(or_pair_number=i)
        if pair_questions.exists():
            part_b_pairs.append((i, pair_questions))
    
    # Calculate distribution
    distribution = {
        'co_distribution': {},
        'bloom_distribution': {},
        'l1_l2_total': 0,
        'l3_l4_total': 0,
        'l5_l6_total': 0,
        'l1_l2_percentage': 0,
        'l3_l4_percentage': 0,
        'l5_l6_percentage': 0,
    }
    
    all_questions = qp.questions.all()
    total_marks = sum(q.marks for q in all_questions)
    
    for co in ['CO1', 'CO2', 'CO3', 'CO4', 'CO5']:
        marks = sum(q.marks for q in all_questions if q.course_outcome == co)
        distribution['co_distribution'][co] = {
            'marks': marks,
            'percentage': (marks / total_marks * 100) if total_marks > 0 else 0
        }
    
    for bl in ['L1', 'L2', 'L3', 'L4', 'L5', 'L6']:
        marks = sum(q.marks for q in all_questions if q.bloom_level == bl)
        distribution['bloom_distribution'][bl] = {
            'marks': marks,
            'percentage': (marks / total_marks * 100) if total_marks > 0 else 0
        }
        
        if bl in ['L1', 'L2']:
            distribution['l1_l2_total'] += marks
        elif bl in ['L3', 'L4']:
            distribution['l3_l4_total'] += marks
        elif bl in ['L5', 'L6']:
            distribution['l5_l6_total'] += marks
    
    if total_marks > 0:
        distribution['l1_l2_percentage'] = distribution['l1_l2_total'] / total_marks * 100
        distribution['l3_l4_percentage'] = distribution['l3_l4_total'] / total_marks * 100
        distribution['l5_l6_percentage'] = distribution['l5_l6_total'] / total_marks * 100
    
    # Validation - use enhanced smart validation
    validation_result = qp.validate_distribution()
    validation_errors = validation_result['errors']
    validation_suggestions = validation_result['suggestions']
    
    # Repetition detection against Question Bank
    repetitions = qp.check_repetitions()
    
    can_submit = len(validation_errors) == 0 and qp.status in ['DRAFT', 'REJECTED']

    auto_checklist_form2 = _build_auto_distribution_checklist(qp)
    checklist_form1 = qp.submission_checklist or {}
    if not checklist_form1:
        checklist_form1 = {
            'faculty_name': qp.faculty.user.full_name,
            'course_code_title': f"{qp.course.course_code} - {qp.course.title}",
            'regulation': getattr(qp.regulation, 'name', ''),
            'month_year': qp.exam_month_year,
            'branch': getattr(getattr(qp.course, 'department', None), 'name', ''),
            'table_charts_list': '',
            'q1': '',
            'q2': '',
            'q3': '',
            'q4': '',
            'q5': '',
            'q6': '',
            'q7': '',
        }
    
    context = {
        'qp': qp,
        'part_a_questions': part_a_questions,
        'part_b_pairs': part_b_pairs,
        'part_c_questions': part_c_questions,
        'distribution': distribution,
        'validation_errors': validation_errors,
        'validation_suggestions': validation_suggestions,
        'repetitions': repetitions,
        'can_submit': can_submit,
        'checklist_form1': checklist_form1,
        'auto_checklist_form2': auto_checklist_form2,
        'page_title': 'Preview Question Paper'
    }
    return render(request, "staff_template/preview_structured_qp.html", context)


@login_required
def staff_submit_structured_qp(request, qp_id):
    """Submit question paper and answer key for HOD review"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied.")
        return redirect('/')
    
    if request.method != 'POST':
        return redirect('staff_preview_structured_qp', qp_id=qp_id)
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    qp = get_object_or_404(StructuredQuestionPaper, id=qp_id, faculty=faculty)
    
    if qp.status not in ['DRAFT', 'REJECTED']:
        messages.warning(request, "Question paper already submitted.")
        return redirect('staff_preview_structured_qp', qp_id=qp.id)

    # For manually uploaded QPs, ensure uploaded source document is available.
    if qp.is_uploaded and not qp.uploaded_document:
        messages.error(request, "Uploaded question paper file is missing. Please upload the QP again before submitting.")
        return redirect('staff_preview_structured_qp', qp_id=qp.id)
    
    # Require answer key document
    answer_key_file = request.FILES.get('answer_key_document')
    if not answer_key_file:
        messages.error(request, "Please upload the answer key document before submitting.")
        return redirect('staff_preview_structured_qp', qp_id=qp.id)

    # Mandatory Form 1 checklist inputs
    required_text_fields = {
        'faculty_name': request.POST.get('checklist_faculty_name', '').strip(),
        'course_code_title': request.POST.get('checklist_course_code_title', '').strip(),
        'regulation': request.POST.get('checklist_regulation', '').strip(),
        'month_year': request.POST.get('checklist_month_year', '').strip(),
        'branch': request.POST.get('checklist_branch', '').strip(),
    }
    checklist_answers = {
        'q1': request.POST.get('checklist_q1', '').strip().lower(),
        'q2': request.POST.get('checklist_q2', '').strip().lower(),
        'q3': request.POST.get('checklist_q3', '').strip().lower(),
        'q4': request.POST.get('checklist_q4', '').strip().lower(),
        'q5': request.POST.get('checklist_q5', '').strip().lower(),
        'q6': request.POST.get('checklist_q6', '').strip().lower(),
        'q7': request.POST.get('checklist_q7', '').strip().lower(),
    }
    table_charts_list = request.POST.get('checklist_table_charts_list', '').strip()

    missing_text = [k for k, v in required_text_fields.items() if not v]
    invalid_answers = [k for k, v in checklist_answers.items() if v not in ['yes', 'no']]
    if missing_text or invalid_answers:
        messages.error(request, "Please complete Checklist Form 1 before submitting for review.")
        return redirect('staff_preview_structured_qp', qp_id=qp.id)

    manual_checklist = {
        **required_text_fields,
        **checklist_answers,
        'table_charts_list': table_charts_list,
    }

    # Auto-filled Form 2 snapshot at submission time
    qp.submission_checklist = manual_checklist
    qp.auto_distribution_checklist = _build_auto_distribution_checklist(qp)
    qp.checklist_completed_at = timezone.now()
    
    # Validate file type
    import os
    ext = os.path.splitext(answer_key_file.name)[1].lower()
    if ext not in ['.pdf', '.doc', '.docx']:
        messages.error(request, "Invalid file type. Allowed: PDF, DOC, DOCX")
        return redirect('staff_preview_structured_qp', qp_id=qp.id)
    
    # Increment revision number on resubmission after REJECTED
    if qp.status == 'REJECTED':
        qp.revision_number += 1
    
    # Save answer key
    if qp.answer_key_document:
        qp.answer_key_document.delete(save=False)
    qp.answer_key_document = answer_key_file
    qp.answer_key_status = 'SUBMITTED'
    qp.answer_key_submitted_at = timezone.now()
    
    qp.status = 'SUBMITTED'
    qp.submitted_at = timezone.now()
    qp.hod_comments = ''  # Clear previous rejection comments
    qp.save()
    
    # Update linked assignment status if exists
    if qp.qp_assignment:
        qp.qp_assignment.status = 'SUBMITTED'
        qp.qp_assignment.save()
    
    messages.success(request, "Question paper and answer key submitted successfully!")
    return redirect('staff_list_structured_qps')


@login_required
def staff_download_structured_qp(request, qp_id):
    """Download generated question paper"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    qp = get_object_or_404(StructuredQuestionPaper, id=qp_id, faculty=faculty)
    
    # Get questions by part
    part_a_questions = qp.questions.filter(part='A').order_by('question_number')
    part_b_questions = qp.questions.filter(part='B').order_by('or_pair_number', 'option_label')
    part_c_questions = qp.questions.filter(part='C')
    
    # Group Part B into OR pairs
    part_b_pairs = []
    for i in range(11, 16):
        pair_questions = list(part_b_questions.filter(or_pair_number=i))
        if pair_questions:
            part_b_pairs.append((i, pair_questions))
    
    # Generate document only when neither generated nor uploaded source is available.
    if not qp.generated_document and not (qp.is_uploaded and qp.uploaded_document):
        from docx import Document
        from docx.shared import Inches, Pt, Twips, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn, nsmap
        from docx.oxml import OxmlElement
        import io
        from django.core.files.base import ContentFile
        
        def set_cell_width(cell, width_inches):
            """Set exact cell width using XML"""
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(width_inches * 1440)))  # Convert inches to twips
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
        
        def set_table_fixed_layout(table, col_widths):
            """Set table to fixed layout with exact column widths"""
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            # Set table layout to fixed
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'fixed')
            tblPr.append(tblLayout)
            
            # Set total table width
            total_width = sum(col_widths)
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), str(int(total_width * 1440)))
            tblW.set(qn('w:type'), 'dxa')
            tblPr.append(tblW)
            
            # Add borders
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border_el = OxmlElement(f'w:{border_name}')
                border_el.set(qn('w:val'), 'single')
                border_el.set(qn('w:sz'), '4')
                border_el.set(qn('w:space'), '0')
                border_el.set(qn('w:color'), '000000')
                tblBorders.append(border_el)
            tblPr.append(tblBorders)
            
            # Set column widths via tblGrid
            tblGrid = OxmlElement('w:tblGrid')
            for width in col_widths:
                gridCol = OxmlElement('w:gridCol')
                gridCol.set(qn('w:w'), str(int(width * 1440)))
                tblGrid.append(gridCol)
            tbl.insert(1, tblGrid)
            
            # Set cell widths for all rows
            for row in table.rows:
                for idx, cell in enumerate(row.cells):
                    if idx < len(col_widths):
                        set_cell_width(cell, col_widths[idx])
        
        def set_cell_text(cell, text, bold=False, size=11, center=False):
            """Set cell text with formatting"""
            cell.text = text
            if cell.paragraphs[0].runs:
                run = cell.paragraphs[0].runs[0]
                run.font.bold = bold
                run.font.size = Pt(size)
                run.font.name = 'Arial'
            if center:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # --- LaTeX-to-OMML helpers for Word math rendering ---
        _xsl_transform_cache = [None]  # mutable container for caching
        
        def _get_xsl_transform():
            """Load and cache the MML2OMML XSLT transform"""
            if _xsl_transform_cache[0] is not None:
                return _xsl_transform_cache[0]
            import os as _os
            from lxml import etree as _et
            for xp in [
                r'C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL',
                r'C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL',
            ]:
                if _os.path.exists(xp):
                    _xsl_transform_cache[0] = _et.XSLT(_et.parse(xp))
                    return _xsl_transform_cache[0]
            return None
        
        def _fix_mathml_matrices(mml_root):
            """Wrap mtable + adjacent bracket <mo> elements in <mrow> with fence attrs
            so MML2OMML.XSL generates proper <m:d> delimiters with stretchy brackets."""
            from lxml import etree as _et
            NS = 'http://www.w3.org/1998/Math/MathML'
            M = f'{{{NS}}}'
            bracket_map = {'(': ')', '[': ']', '{': '}'}
            for mtable in list(mml_root.iter(f'{M}mtable')):
                parent = mtable.getparent()
                if parent is None:
                    continue
                children = list(parent)
                idx = children.index(mtable)
                open_mo = close_mo = None
                if idx > 0:
                    prev = children[idx - 1]
                    if prev.tag == f'{M}mo' and (prev.text or '').strip() in bracket_map:
                        open_mo = prev
                if idx < len(children) - 1:
                    nxt = children[idx + 1]
                    if nxt.tag == f'{M}mo' and (nxt.text or '').strip() in bracket_map.values():
                        close_mo = nxt
                if open_mo is not None and close_mo is not None:
                    for mo in (open_mo, close_mo):
                        mo.set('stretchy', 'true')
                        mo.set('fence', 'true')
                    open_mo.set('form', 'prefix')
                    close_mo.set('form', 'postfix')
                    wrapper = _et.Element(f'{M}mrow')
                    parent.insert(children.index(open_mo), wrapper)
                    parent.remove(open_mo)
                    parent.remove(mtable)
                    parent.remove(close_mo)
                    wrapper.append(open_mo)
                    wrapper.append(mtable)
                    wrapper.append(close_mo)
            return mml_root
        
        def _fix_omml_separators(omml_root):
            """Replace OMML <m:d> elements that use comma separators with flattened
            explicit comma runs. Word renders <m:sepChr val=","> as middle dots;
            this fix makes commas display as actual commas."""
            from lxml import etree as _et
            NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
            M = f'{{{NS}}}'
            for d_elem in list(omml_root.iter(f'{M}d')):
                dPr = d_elem.find(f'{M}dPr')
                if dPr is None:
                    continue
                sep_el = dPr.find(f'{M}sepChr')
                if sep_el is None:
                    continue
                sep_char = sep_el.get(f'{M}val', ',')
                if sep_char != ',':
                    continue
                beg_el = dPr.find(f'{M}begChr')
                end_el = dPr.find(f'{M}endChr')
                beg_char = beg_el.get(f'{M}val', '(') if beg_el is not None else '('
                end_char = end_el.get(f'{M}val', ')') if end_el is not None else ')'
                elements = list(d_elem.findall(f'{M}e'))
                parent = d_elem.getparent()
                idx = list(parent).index(d_elem)
                def _mk(text):
                    r = _et.Element(f'{M}r')
                    t = _et.SubElement(r, f'{M}t')
                    t.text = text
                    return r
                parts = [_mk(beg_char)]
                for i, e in enumerate(elements):
                    for child in list(e):
                        parts.append(child)
                    if i < len(elements) - 1:
                        parts.append(_mk(', '))
                parts.append(_mk(end_char))
                for j, p in enumerate(parts):
                    parent.insert(idx + j, p)
                parent.remove(d_elem)
            return omml_root
        
        def _latex_to_omml(latex_str):
            """Convert a LaTeX string to an OMML XML element for python-docx.
            Returns the element, or None on failure."""
            try:
                import latex2mathml.converter
                from lxml import etree as _et
                xsl = _get_xsl_transform()
                if xsl is None:
                    return None
                mml = latex2mathml.converter.convert(latex_str)
                mml_tree = _et.fromstring(mml.encode())
                _fix_mathml_matrices(mml_tree)
                omml_tree = xsl(mml_tree)
                omml_root = omml_tree.getroot()
                _fix_omml_separators(omml_root)
                return omml_root
            except Exception:
                return None
        
        def add_question_content(cell, question, size=11, center=False):
            """Set cell text with LaTeX math rendered as native Word equations,
            and optionally embed question image below it."""
            import re as _re
            
            text = question.question_text or ''
            
            # Split on $$...$$ (block) and $...$ (inline) math delimiters
            pattern = _re.compile(r'(\$\$.*?\$\$|\$.*?\$)', _re.DOTALL)
            segments = pattern.split(text)
            has_math = any(s.startswith('$') for s in segments if s)
            
            # Clear cell and get primary paragraph
            cell.text = ''
            para = cell.paragraphs[0]
            if center:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if not has_math:
                # No math delimiters — plain text
                run = para.add_run(text)
                run.font.size = Pt(size)
                run.font.name = 'Arial'
            else:
                for segment in segments:
                    if not segment:
                        continue
                    
                    is_block = segment.startswith('$$') and segment.endswith('$$')
                    is_inline = (not is_block) and segment.startswith('$') and segment.endswith('$')
                    
                    if is_block or is_inline:
                        latex = segment[2:-2].strip() if is_block else segment[1:-1].strip()
                        
                        if is_block:
                            target_para = cell.add_paragraph()
                            target_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            target_para = para
                        
                        omml_el = _latex_to_omml(latex)
                        if omml_el is not None:
                            target_para._p.append(omml_el)
                        else:
                            # Fallback: raw LaTeX as italic
                            r = target_para.add_run(segment)
                            r.font.size = Pt(size)
                            r.font.name = 'Arial'
                            r.italic = True
                    else:
                        # Plain text — handle newlines
                        lines = segment.split('\n')
                        for li, line in enumerate(lines):
                            if line:
                                r = para.add_run(line)
                                r.font.size = Pt(size)
                                r.font.name = 'Arial'
                            if li < len(lines) - 1:
                                r = para.add_run()
                                r.add_break()
            
            # Embed question image if available
            if question.question_image:
                try:
                    img_para = cell.add_paragraph()
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_run = img_para.add_run()
                    img_run.add_picture(question.question_image.path, width=Inches(3.5))
                except Exception:
                    pass  # Skip image if file not found or invalid
        
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
        
        # Roll No - Right aligned with actual table boxes
        rollno_table = doc.add_table(rows=1, cols=11)
        rollno_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        
        # Set "Roll No." text in first cell (no border, wider width, bottom aligned)
        rollno_cell = rollno_table.rows[0].cells[0]
        rollno_cell.text = "Roll No."
        rollno_cell.paragraphs[0].runs[0].font.bold = True
        rollno_cell.paragraphs[0].runs[0].font.size = Pt(12)
        rollno_cell.paragraphs[0].runs[0].font.name = 'Arial'
        rollno_cell.width = Inches(0.9)
        # Set vertical alignment to bottom
        tc_rollno = rollno_cell._tc
        tcPr_rollno = tc_rollno.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'bottom')
        tcPr_rollno.append(vAlign)
        # Set fixed width for Roll No. cell
        tcW_rollno = OxmlElement('w:tcW')
        tcW_rollno.set(qn('w:w'), str(int(0.9 * 1440)))
        tcW_rollno.set(qn('w:type'), 'dxa')
        tcPr_rollno.append(tcW_rollno)
        
        # Add bordered boxes for each digit
        for i in range(1, 11):
            box_cell = rollno_table.rows[0].cells[i]
            box_cell.width = Inches(0.28)
            # Set cell borders
            tc = box_cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border_el = OxmlElement(f'w:{border_name}')
                border_el.set(qn('w:val'), 'single')
                border_el.set(qn('w:sz'), '4')
                border_el.set(qn('w:space'), '0')
                border_el.set(qn('w:color'), '000000')
                tcBorders.append(border_el)
            tcPr.append(tcBorders)
            # Set fixed width
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(0.28 * 1440)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
        
        doc.add_paragraph()  # spacing after roll no
        
        # University Header
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para.paragraph_format.space_before = Pt(0)
        header_para.paragraph_format.space_after = Pt(6)
        header_run = header_para.add_run("ANNA UNIVERSITY (UNIVERSITY DEPARTMENTS)")
        header_run.font.bold = True
        header_run.font.size = Pt(13)
        header_run.font.name = 'Arial'
        
        # Exam Info
        exam_para = doc.add_paragraph()
        exam_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        exam_para.paragraph_format.space_after = Pt(6)
        if qp.qp_assignment:
            exam_type_label = qp.qp_assignment.get_exam_type_display().upper()
        else:
            exam_type_label = getattr(qp.semester, 'semester_type', 'END SEMESTER') + ' SEMESTER'
        exam_run = exam_para.add_run(f"B.E. / B. Tech (Full Time) - {exam_type_label} EXAMINATIONS, {qp.exam_month_year.upper()}")
        exam_run.font.bold = True
        exam_run.font.size = Pt(12)
        exam_run.font.name = 'Arial'
        
        # Course Title
        course_title_para = doc.add_paragraph()
        course_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        course_title_para.paragraph_format.space_after = Pt(3)
        title_run = course_title_para.add_run(f"{qp.course.title.upper()}")
        title_run.font.bold = True
        title_run.font.size = Pt(12)
        title_run.font.name = 'Arial'
        
        # Semester info
        sem_para = doc.add_paragraph()
        sem_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sem_para.paragraph_format.space_after = Pt(3)
        sem_num = getattr(qp.semester, 'semester_number', '')
        sem_run = sem_para.add_run(f"Semester {sem_num}" if sem_num else "Semester")
        sem_run.font.size = Pt(11)
        sem_run.font.name = 'Arial'
        
        # Course Code and Name
        code_para = doc.add_paragraph()
        code_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        code_para.paragraph_format.space_after = Pt(3)
        code_run = code_para.add_run(f"{qp.course.course_code} - {qp.course.title}")
        code_run.font.bold = True
        code_run.font.size = Pt(12)
        code_run.font.name = 'Arial'
        
        # Regulation
        reg_para = doc.add_paragraph()
        reg_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        reg_para.paragraph_format.space_after = Pt(12)
        reg_run = reg_para.add_run(f"(Regulation {qp.regulation.name})")
        reg_run.font.size = Pt(11)
        reg_run.font.name = 'Arial'
        
        # Time and Marks - use table for proper alignment
        time_marks_table = doc.add_table(rows=1, cols=2)
        time_marks_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        time_cell = time_marks_table.rows[0].cells[0]
        marks_cell = time_marks_table.rows[0].cells[1]
        
        time_cell.text = "Time: 3 hours"
        time_cell.paragraphs[0].runs[0].font.bold = True
        time_cell.paragraphs[0].runs[0].font.size = Pt(12)
        time_cell.paragraphs[0].runs[0].font.name = 'Arial'
        
        marks_cell.text = "Max. Marks: 100"
        marks_cell.paragraphs[0].runs[0].font.bold = True
        marks_cell.paragraphs[0].runs[0].font.size = Pt(12)
        marks_cell.paragraphs[0].runs[0].font.name = 'Arial'
        marks_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # spacing
        
        # CO Table - 2 columns with descriptions
        co_descriptions = [
            ('CO 1', getattr(qp, 'co1_description', '') or ''),
            ('CO 2', getattr(qp, 'co2_description', '') or ''),
            ('CO 3', getattr(qp, 'co3_description', '') or ''),
            ('CO 4', getattr(qp, 'co4_description', '') or ''),
            ('CO 5', getattr(qp, 'co5_description', '') or ''),
        ]
        
        co_table = doc.add_table(rows=5, cols=2)
        co_col_widths = [0.6, 6.5]  # inches
        set_table_fixed_layout(co_table, co_col_widths)
        
        for i, (label, desc) in enumerate(co_descriptions):
            set_cell_text(co_table.rows[i].cells[0], label, bold=True, size=11, center=True)
            set_cell_text(co_table.rows[i].cells[1], desc, bold=False, size=11, center=False)
        
        # BL Legend
        bl_para = doc.add_paragraph()
        bl_para.paragraph_format.space_before = Pt(6)
        bl_para.paragraph_format.space_after = Pt(12)
        
        bl_label = bl_para.add_run("BL – Bloom's Taxonomy Levels")
        bl_label.font.bold = True
        bl_label.font.size = Pt(11)
        bl_label.font.name = 'Arial'
        
        bl_para.add_run("\n")
        bl_desc = bl_para.add_run("(L1 - Remembering, L2 - Understanding, L3 - Applying, L4 - Analysing, L5 - Evaluating, L6 - Creating)")
        bl_desc.font.size = Pt(10)
        bl_desc.font.name = 'Arial'
        
        # ============ PART A ============
        part_a_title = doc.add_paragraph()
        part_a_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part_a_title.paragraph_format.space_before = Pt(12)
        part_a_title.paragraph_format.space_after = Pt(3)
        a_run = part_a_title.add_run("PART- A (10 x 2 = 20 Marks)")
        a_run.font.bold = True
        a_run.font.size = Pt(12)
        a_run.font.name = 'Arial'
        a_run.font.underline = True
        
        part_a_subtitle = doc.add_paragraph()
        part_a_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part_a_subtitle.paragraph_format.space_after = Pt(6)
        sub_run = part_a_subtitle.add_run("(Answer all Questions)")
        sub_run.font.size = Pt(11)
        sub_run.font.name = 'Arial'
        
        # Part A table - FIXED column widths
        part_a_col_widths = [0.55, 5.0, 0.55, 0.5, 0.5]  # Q.No, Questions, Marks, CO, BL
        part_a_table = doc.add_table(rows=11, cols=5)
        set_table_fixed_layout(part_a_table, part_a_col_widths)
        
        # Headers
        headers = ['Q. No', 'Questions', 'Marks', 'CO', 'BL']
        for i, header in enumerate(headers):
            set_cell_text(part_a_table.rows[0].cells[i], header, bold=True, size=11, center=True)
        
        # Questions
        for idx, q in enumerate(part_a_questions[:10], 1):
            row = part_a_table.rows[idx]
            set_cell_text(row.cells[0], str(q.question_number), size=11, center=True)
            add_question_content(row.cells[1], q, size=11, center=False)
            set_cell_text(row.cells[2], str(q.marks), size=11, center=True)
            # Extract just the number from CO1, CO2, etc.
            co_num = (q.course_outcome or '').replace('CO', '').strip()
            set_cell_text(row.cells[3], co_num, size=11, center=True)
            set_cell_text(row.cells[4], q.bloom_level or '', size=11, center=True)
        
        # ============ PART B ============
        part_b_title = doc.add_paragraph()
        part_b_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part_b_title.paragraph_format.space_before = Pt(12)
        part_b_title.paragraph_format.space_after = Pt(3)
        b_run = part_b_title.add_run("PART- B (5 x 13 = 65 Marks)")
        b_run.font.bold = True
        b_run.font.size = Pt(12)
        b_run.font.name = 'Arial'
        b_run.font.underline = True
        
        part_b_subtitle = doc.add_paragraph()
        part_b_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part_b_subtitle.paragraph_format.space_after = Pt(6)
        sub_run = part_b_subtitle.add_run("(Answer any FIVE questions, choosing one from each unit)")
        sub_run.font.size = Pt(11)
        sub_run.font.name = 'Arial'
        
        # Part B table
        part_b_row_count = 1 + len(part_b_pairs) * 3  # header + (a, OR, b) for each pair
        part_b_table = doc.add_table(rows=1, cols=5)
        part_b_col_widths = [0.55, 5.0, 0.55, 0.5, 0.5]
        set_table_fixed_layout(part_b_table, part_b_col_widths)
        
        # Headers
        for i, header in enumerate(headers):
            set_cell_text(part_b_table.rows[0].cells[i], header, bold=True, size=11, center=True)
        
        # Questions
        for pair_num, questions in part_b_pairs:
            # Option (a)
            row_a = part_b_table.add_row()
            set_cell_width(row_a.cells[0], part_b_col_widths[0])
            set_cell_width(row_a.cells[1], part_b_col_widths[1])
            set_cell_width(row_a.cells[2], part_b_col_widths[2])
            set_cell_width(row_a.cells[3], part_b_col_widths[3])
            set_cell_width(row_a.cells[4], part_b_col_widths[4])
            
            set_cell_text(row_a.cells[0], f"{pair_num} (a)", size=11, center=True)
            q_a = next((q for q in questions if q.option_label == '(a)'), None)
            if q_a:
                add_question_content(row_a.cells[1], q_a, size=11, center=False)
                co_num = (q_a.course_outcome or '').replace('CO', '').strip()
                set_cell_text(row_a.cells[3], co_num, size=11, center=True)
                set_cell_text(row_a.cells[4], q_a.bloom_level or '', size=11, center=True)
            set_cell_text(row_a.cells[2], '13', size=11, center=True)
            
            # OR row
            or_row = part_b_table.add_row()
            set_cell_width(or_row.cells[0], part_b_col_widths[0])
            set_cell_width(or_row.cells[1], part_b_col_widths[1])
            set_cell_width(or_row.cells[2], part_b_col_widths[2])
            set_cell_width(or_row.cells[3], part_b_col_widths[3])
            set_cell_width(or_row.cells[4], part_b_col_widths[4])
            set_cell_text(or_row.cells[0], '', size=11, center=True)
            set_cell_text(or_row.cells[1], 'OR', bold=True, size=11, center=True)
            set_cell_text(or_row.cells[2], '', size=11, center=True)
            set_cell_text(or_row.cells[3], '', size=11, center=True)
            set_cell_text(or_row.cells[4], '', size=11, center=True)
            
            # Option (b)
            row_b = part_b_table.add_row()
            set_cell_width(row_b.cells[0], part_b_col_widths[0])
            set_cell_width(row_b.cells[1], part_b_col_widths[1])
            set_cell_width(row_b.cells[2], part_b_col_widths[2])
            set_cell_width(row_b.cells[3], part_b_col_widths[3])
            set_cell_width(row_b.cells[4], part_b_col_widths[4])
            
            set_cell_text(row_b.cells[0], f"{pair_num} (b)", size=11, center=True)
            q_b = next((q for q in questions if q.option_label == '(b)'), None)
            if q_b:
                add_question_content(row_b.cells[1], q_b, size=11, center=False)
                co_num = (q_b.course_outcome or '').replace('CO', '').strip()
                set_cell_text(row_b.cells[3], co_num, size=11, center=True)
                set_cell_text(row_b.cells[4], q_b.bloom_level or '', size=11, center=True)
            set_cell_text(row_b.cells[2], '13', size=11, center=True)
        
        # ============ PART C ============
        part_c_title = doc.add_paragraph()
        part_c_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part_c_title.paragraph_format.space_before = Pt(12)
        part_c_title.paragraph_format.space_after = Pt(3)
        c_run = part_c_title.add_run("PART- C (1 x 15 = 15 Marks)")
        c_run.font.bold = True
        c_run.font.size = Pt(12)
        c_run.font.name = 'Arial'
        c_run.font.underline = True
        
        part_c_subtitle = doc.add_paragraph()
        part_c_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part_c_subtitle.paragraph_format.space_after = Pt(6)
        sub_run = part_c_subtitle.add_run("(Q.No. 16 is compulsory)")
        sub_run.font.size = Pt(11)
        sub_run.font.name = 'Arial'
        
        # Part C table
        part_c_table = doc.add_table(rows=2, cols=5)
        part_c_col_widths = [0.55, 5.0, 0.55, 0.5, 0.5]
        set_table_fixed_layout(part_c_table, part_c_col_widths)
        
        # Headers
        for i, header in enumerate(headers):
            set_cell_text(part_c_table.rows[0].cells[i], header, bold=True, size=11, center=True)
        
        # Question
        if part_c_questions.exists():
            q = part_c_questions.first()
            row = part_c_table.rows[1]
            set_cell_text(row.cells[0], '16', size=11, center=True)
            add_question_content(row.cells[1], q, size=11, center=False)
            set_cell_text(row.cells[2], str(q.marks), size=11, center=True)
            co_num = (q.course_outcome or '').replace('CO', '').strip()
            set_cell_text(row.cells[3], co_num, size=11, center=True)
            set_cell_text(row.cells[4], q.bloom_level or '', size=11, center=True)
        
        # Save to ContentFile
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        qp.generated_document.save(f'qp_{qp.id}.docx', ContentFile(doc_io.read()), save=True)
    
    from django.http import FileResponse

    base_field = qp.generated_document if qp.generated_document else qp.uploaded_document
    if not base_field:
        messages.error(request, "Question paper document is not available for download.")
        return redirect('staff_preview_structured_qp', qp_id=qp.id)

    # Build one complete DOCX by appending checklist forms at the end.
    checklist_exists = bool(qp.submission_checklist or qp.auto_distribution_checklist)
    source_name = base_field.name.split('/')[-1]
    is_docx = source_name.lower().endswith('.docx')
    if checklist_exists and is_docx:
        from docx import Document
        import io

        with base_field.open('rb') as fp:
            doc = Document(fp)

        append_checklist_to_document(doc, qp)

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        response = HttpResponse(
            output.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="qp_{qp.id}_complete.docx"'
        return response

    return FileResponse(base_field.open('rb'), as_attachment=True, filename=source_name)


@login_required
def staff_manage_qp_answers(request, qp_id):
    """Page for managing answers for each question using AI generation"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    qp = get_object_or_404(StructuredQuestionPaper, id=qp_id, faculty=faculty)
    
    # Get questions by part
    part_a_questions = qp.questions.filter(part='A').order_by('question_number')
    part_b_questions = qp.questions.filter(part='B').order_by('question_number', 'option_label')
    part_c_questions = qp.questions.filter(part='C').order_by('question_number')
    
    # Group Part B questions by OR pairs
    part_b_pairs = []
    for pair_num in [11, 12, 13, 14, 15]:
        pair_questions = part_b_questions.filter(or_pair_number=pair_num).order_by('option_label')
        if pair_questions.exists():
            part_b_pairs.append((pair_num, list(pair_questions)))
    
    # Count answered questions
    all_questions = qp.questions.all()
    answered_count = all_questions.exclude(answer='').exclude(answer__isnull=True).count()
    total_questions = all_questions.count()
    
    context = {
        'page_title': f'Manage Answers - {qp.course.course_code}',
        'qp': qp,
        'part_a_questions': part_a_questions,
        'part_b_pairs': part_b_pairs,
        'part_c_questions': part_c_questions,
        'answered_count': answered_count,
        'total_questions': total_questions,
    }
    
    return render(request, 'staff_template/manage_qp_answers.html', context)


@login_required
@csrf_exempt
def staff_generate_answer_options(request):
    """AJAX endpoint to generate AI answer options for a question"""
    if not check_faculty_permission(request.user):
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    if request.method != 'POST':
        return JsonResponse({'error': 'POST method required'}, status=405)
    
    try:
        data = json.loads(request.body)
        question_text = data.get('question', '').strip()
        marks = int(data.get('marks', 2))
        part_type = data.get('part', 'A')
        course_name = data.get('course_name', '')
        
        if not question_text:
            return JsonResponse({'error': 'Question text is required'}, status=400)
        
        from .utils.ai_answer_generator import generate_answer_options
        
        answers = generate_answer_options(
            question_text=question_text,
            marks=marks,
            part_type=part_type,
            course_name=course_name,
            num_options=4
        )
        
        return JsonResponse({'answers': answers})
        
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
def staff_download_answer_key(request, qp_id):
    """Download answer key for a structured question paper"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    qp = get_object_or_404(StructuredQuestionPaper, id=qp_id, faculty=faculty)
    
    # Check if any answers exist
    questions_with_answers = qp.questions.exclude(answer='').exclude(answer__isnull=True)
    if not questions_with_answers.exists():
        messages.warning(request, "No answers have been added to this question paper yet.")
        return redirect('staff_preview_structured_qp', qp_id=qp.id)
    
    # Generate answer key document
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io
    
    doc = Document()
    
    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("ANSWER KEY")
    title_run.font.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = 'Arial'
    
    # Course details
    details_para = doc.add_paragraph()
    details_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details_run = details_para.add_run(f"{qp.course.course_code} - {qp.course.title}")
    details_run.font.bold = True
    details_run.font.size = Pt(12)
    details_run.font.name = 'Arial'
    
    exam_para = doc.add_paragraph()
    exam_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    exam_run = exam_para.add_run(f"Examination: {qp.exam_month_year}")
    exam_run.font.size = Pt(11)
    exam_run.font.name = 'Arial'
    
    doc.add_paragraph()  # Spacing
    
    # Part A
    part_a_questions = qp.questions.filter(part='A').order_by('question_number')
    if part_a_questions.exists():
        part_a_title = doc.add_paragraph()
        part_a_run = part_a_title.add_run("PART A - Short Answers (2 marks each)")
        part_a_run.font.bold = True
        part_a_run.font.size = Pt(12)
        part_a_run.font.name = 'Arial'
        part_a_run.font.underline = True
        
        for q in part_a_questions:
            # Question
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"Q{q.question_number}. {q.question_text}")
            q_run.font.bold = True
            q_run.font.size = Pt(11)
            q_run.font.name = 'Arial'
            
            # Answer
            if q.answer:
                a_para = doc.add_paragraph()
                a_run = a_para.add_run(f"Answer: {q.answer}")
                a_run.font.size = Pt(11)
                a_run.font.name = 'Arial'
            else:
                a_para = doc.add_paragraph()
                a_run = a_para.add_run("Answer: [Not provided]")
                a_run.font.size = Pt(11)
                a_run.font.name = 'Arial'
                a_run.font.italic = True
            
            doc.add_paragraph()  # Spacing
    
    # Part B
    part_b_questions = qp.questions.filter(part='B').order_by('or_pair_number', 'option_label')
    if part_b_questions.exists():
        part_b_title = doc.add_paragraph()
        part_b_run = part_b_title.add_run("PART B - Descriptive Answers (13 marks each)")
        part_b_run.font.bold = True
        part_b_run.font.size = Pt(12)
        part_b_run.font.name = 'Arial'
        part_b_run.font.underline = True
        
        for q in part_b_questions:
            # Question
            q_label = f"Q{q.or_pair_number}{q.option_label}" if q.option_label else f"Q{q.question_number}"
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"{q_label}. {q.question_text}")
            q_run.font.bold = True
            q_run.font.size = Pt(11)
            q_run.font.name = 'Arial'
            
            # Answer
            if q.answer:
                a_para = doc.add_paragraph()
                a_run = a_para.add_run(f"Answer: {q.answer}")
                a_run.font.size = Pt(11)
                a_run.font.name = 'Arial'
            else:
                a_para = doc.add_paragraph()
                a_run = a_para.add_run("Answer: [Not provided]")
                a_run.font.size = Pt(11)
                a_run.font.name = 'Arial'
                a_run.font.italic = True
            
            doc.add_paragraph()  # Spacing
    
    # Part C
    part_c_questions = qp.questions.filter(part='C')
    if part_c_questions.exists():
        part_c_title = doc.add_paragraph()
        part_c_run = part_c_title.add_run("PART C - Problem Solving (15 marks)")
        part_c_run.font.bold = True
        part_c_run.font.size = Pt(12)
        part_c_run.font.name = 'Arial'
        part_c_run.font.underline = True
        
        for q in part_c_questions:
            # Question
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"Q16. {q.question_text}")
            q_run.font.bold = True
            q_run.font.size = Pt(11)
            q_run.font.name = 'Arial'
            
            # Answer
            if q.answer:
                a_para = doc.add_paragraph()
                a_run = a_para.add_run(f"Answer: {q.answer}")
                a_run.font.size = Pt(11)
                a_run.font.name = 'Arial'
            else:
                a_para = doc.add_paragraph()
                a_run = a_para.add_run("Answer: [Not provided]")
                a_run.font.size = Pt(11)
                a_run.font.name = 'Arial'
                a_run.font.italic = True
    
    # Save to response
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    from django.http import FileResponse
    filename = f'Answer_Key_{qp.course.course_code}_{qp.exam_month_year.replace("/", "_")}.docx'
    return FileResponse(doc_io, as_attachment=True, filename=filename)


@login_required
@csrf_exempt
def staff_save_question_answer(request):
    """AJAX endpoint to save an answer for a question"""
    if not check_faculty_permission(request.user):
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    if request.method != 'POST':
        return JsonResponse({'error': 'POST method required'}, status=405)
    
    try:
        data = json.loads(request.body)
        question_id = data.get('question_id')
        answer = data.get('answer', '').strip()
        
        if not question_id:
            return JsonResponse({'error': 'Question ID is required'}, status=400)
        
        faculty = get_object_or_404(Faculty_Profile, user=request.user)
        question = get_object_or_404(QPQuestion, id=question_id, question_paper__faculty=faculty)
        
        question.answer = answer
        question.save()
        
        return JsonResponse({'success': True, 'message': 'Answer saved successfully'})
        
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
def staff_delete_qp_question(request, question_id):
    """Delete a question from a question paper (only for DRAFT status)"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    question = get_object_or_404(QPQuestion, id=question_id, question_paper__faculty=faculty)
    
    qp = question.question_paper
    
    # Only allow deletion for draft or rejected QPs
    if qp.status not in ['DRAFT', 'REJECTED']:
        messages.error(request, "Cannot delete questions from a submitted question paper.")
        return redirect('staff_preview_structured_qp', qp_id=qp.id)
    
    # Store info for message
    part = question.part
    q_num = question.question_number
    
    # Delete the question
    question.delete()
    
    messages.success(request, f"Question {q_num} from Part {part} deleted successfully.")
    
    # Redirect back to the referring page or preview
    referer = request.META.get('HTTP_REFERER')
    if referer:
        return redirect(referer)
    return redirect('staff_preview_structured_qp', qp_id=qp.id)


@login_required
@csrf_exempt
def staff_delete_qp_question_ajax(request):
    """AJAX endpoint to delete a question from a question paper"""
    if not check_faculty_permission(request.user):
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    if request.method != 'POST':
        return JsonResponse({'error': 'POST method required'}, status=405)
    
    try:
        data = json.loads(request.body)
        question_id = data.get('question_id')
        
        if not question_id:
            return JsonResponse({'error': 'Question ID is required'}, status=400)
        
        faculty = get_object_or_404(Faculty_Profile, user=request.user)
        question = get_object_or_404(QPQuestion, id=question_id, question_paper__faculty=faculty)
        
        qp = question.question_paper
        
        # Only allow deletion for draft or rejected QPs
        if qp.status not in ['DRAFT', 'REJECTED']:
            return JsonResponse({'error': 'Cannot delete questions from a submitted question paper.'}, status=403)
        
        # Store info for response
        part = question.part
        q_num = question.question_number
        
        # Delete the question
        question.delete()
        
        return JsonResponse({
            'success': True, 
            'message': f'Question {q_num} from Part {part} deleted successfully.'
        })
        
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
def staff_delete_structured_qp(request, qp_id):
    """Delete an entire structured question paper (only for DRAFT status)"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied.")
        return redirect('/')
    
    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    qp = get_object_or_404(StructuredQuestionPaper, id=qp_id, faculty=faculty)
    
    # Only allow deletion for draft or rejected QPs
    if qp.status not in ['DRAFT', 'REJECTED']:
        messages.error(request, "Cannot delete a submitted or approved question paper. Only draft or rejected question papers can be deleted.")
        return redirect('staff_list_structured_qps')
    
    if request.method == 'POST':
        # Store info for message
        course_code = qp.course.course_code
        exam_month_year = qp.exam_month_year
        
        # Delete all associated questions first, then the QP
        qp.questions.all().delete()
        qp.delete()
        
        messages.success(request, f"Question paper for {course_code} ({exam_month_year}) deleted successfully.")
        return redirect('staff_list_structured_qps')
    
    # GET request - show confirmation page
    context = {
        'qp': qp,
        'page_title': 'Delete Question Paper'
    }
    return render(request, "staff_template/delete_structured_qp_confirm.html", context)


@login_required
@csrf_exempt
def staff_delete_structured_qp_ajax(request):
    """AJAX endpoint to delete an entire structured question paper"""
    if not check_faculty_permission(request.user):
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    if request.method != 'POST':
        return JsonResponse({'error': 'POST method required'}, status=405)
    
    try:
        data = json.loads(request.body)
        qp_id = data.get('qp_id')
        
        if not qp_id:
            return JsonResponse({'error': 'Question Paper ID is required'}, status=400)
        
        faculty = get_object_or_404(Faculty_Profile, user=request.user)
        qp = get_object_or_404(StructuredQuestionPaper, id=qp_id, faculty=faculty)
        
        # Only allow deletion for draft or rejected QPs
        if qp.status not in ['DRAFT', 'REJECTED']:
            return JsonResponse({'error': 'Cannot delete a submitted or approved question paper. Only draft or rejected question papers can be deleted.'}, status=403)
        
        # Store info for response
        course_code = qp.course.course_code
        exam_month_year = qp.exam_month_year
        
        # Delete all associated questions first, then the QP
        qp.questions.all().delete()
        qp.delete()
        
        return JsonResponse({
            'success': True, 
            'message': f'Question paper for {course_code} ({exam_month_year}) deleted successfully.'
        })
        
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


# =============================================================================
# ANSWER KEY SUBMISSION
# =============================================================================

@login_required
def staff_submit_answer_key(request, qp_id):
    """Faculty uploads answer key document after QP is approved"""
    if not check_faculty_permission(request.user):
        messages.error(request, "Access Denied.")
        return redirect('/')

    faculty = get_object_or_404(Faculty_Profile, user=request.user)
    qp = get_object_or_404(StructuredQuestionPaper, id=qp_id, faculty=faculty)

    if qp.status != 'APPROVED':
        messages.warning(request, "Answer key can only be submitted for approved question papers.")
        return redirect('staff_list_structured_qps')

    if request.method == 'POST':
        answer_key_file = request.FILES.get('answer_key_document')
        if not answer_key_file:
            messages.error(request, "Please upload an answer key document.")
        else:
            # Validate file type
            allowed_types = ['.pdf', '.doc', '.docx']
            import os
            ext = os.path.splitext(answer_key_file.name)[1].lower()
            if ext not in allowed_types:
                messages.error(request, f"Invalid file type. Allowed: {', '.join(allowed_types)}")
            else:
                # Delete old file if exists
                if qp.answer_key_document:
                    qp.answer_key_document.delete(save=False)

                qp.answer_key_document = answer_key_file
                qp.answer_key_status = 'SUBMITTED'
                qp.answer_key_submitted_at = timezone.now()
                qp.save()

                # Notify HOD
                try:
                    from .models import NotificationRecipient
                    hod_profiles = Faculty_Profile.objects.filter(
                        user__user_type=2
                    ).select_related('user')
                    notification = Notification.objects.create(
                        type='answer_key_submitted',
                        message=f"Faculty {faculty.user.first_name} {faculty.user.last_name} has submitted "
                                f"the answer key for {qp.course.course_code} - {qp.course.title} "
                                f"({qp.exam_month_year}). Please review.",
                        created_by=request.user
                    )
                    for hod in hod_profiles:
                        NotificationRecipient.objects.create(
                            notification=notification, recipient=hod.user
                        )
                except Exception:
                    pass  # Non-critical

                messages.success(request, "Answer key submitted successfully! HOD will review it.")
                return redirect('staff_list_structured_qps')

    context = {
        'qp': qp,
        'page_title': f'Submit Answer Key - {qp.course.course_code}'
    }
    return render(request, "staff_template/submit_answer_key.html", context)
