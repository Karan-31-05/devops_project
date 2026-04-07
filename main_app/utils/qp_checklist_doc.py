"""Utilities to generate submission checklist appendix documents for QP workflow."""

import io
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION


FONT_NAME = 'Arial'


def _yn(value):
    value = (value or '').strip().lower()
    if value == 'yes':
        return 'Yes'
    if value == 'no':
        return 'No'
    return '-'


def _style_paragraph(paragraph, size=11, bold=False, align=None, underline=False, before=0, after=2):
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    if not paragraph.runs:
        paragraph.add_run('')
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.underline = underline


def _style_cell(cell, size=10, bold=False, center=False):
    if not cell.paragraphs:
        cell.text = ''
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(size)
        run.font.bold = bold


def _set_table_fixed_widths(table, widths):
    """Set fixed table widths in inches for each column."""
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                cell.width = Inches(widths[idx])


def _set_run_font(run, size=10, bold=False, underline=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.underline = underline


def _get_reference_template_path():
    return Path(__file__).resolve().parent.parent / 'resources' / 'checklist_reference.docx'


def _fill_reference_template(reference_doc, qp, form1, form2):
    """Fill known text/table placeholders in the strict reference template."""
    paras = reference_doc.paragraphs

    label_to_value = {
        'NAME OF THE FACULTY MEMBER': form1.get('faculty_name', ''),
        'COURSE CODE & TITLE': form1.get('course_code_title', ''),
        'REGULATION': form1.get('regulation', ''),
        'MONTH & YEAR': form1.get('month_year', ''),
        'BRANCH': form1.get('branch', ''),
    }

    for para in paras:
        raw = ' '.join((para.text or '').split())
        if not raw:
            continue

        # Fill main declaration fields by matching label text.
        for label, value in label_to_value.items():
            if raw.startswith(label):
                para.text = f"{label} : {value}"
                _style_paragraph(para, size=11, bold=True, after=2)
                break

        # Fill table charts line.
        if raw.startswith('NAME OF THE TABLE CHARTS'):
            charts = form1.get('table_charts_list', '') or '.................................................................'
            para.text = f"NAME OF THE TABLE CHARTS: {charts}"
            _style_paragraph(para, size=11, bold=True, after=2)

    # Table index 5 in this reference holds the 7 checklist yes/no rows.
    if len(reference_doc.tables) > 5:
        checklist_table = reference_doc.tables[5]
        for r, key in enumerate(['q1', 'q2', 'q3', 'q4', 'q5', 'q6', 'q7']):
            if r >= len(checklist_table.rows):
                break
            ans = (form1.get(key, '') or '').strip().lower()
            cell = checklist_table.cell(r, 1)
            if ans == 'yes':
                display_value = '- Yes'
            elif ans == 'no':
                display_value = '- No'
            else:
                display_value = '-'

            cell.text = display_value
            _style_cell(cell, size=11, center=True)

    # Table index 7 in this reference is the working checklist mark-distribution grid.
    if len(reference_doc.tables) > 7:
        grid = reference_doc.tables[7]
        row_map = {int(item.get('question_no')): item for item in form2.get('rows', []) if item.get('question_no')}

        def _fmt_mark(v):
            try:
                iv = int(v)
                return str(iv) if iv > 0 else '-'
            except Exception:
                return '-'

        for q_no in range(1, 17):
            row_idx = q_no + 1  # r2..r17 are Q1..Q16 in template
            if row_idx >= len(grid.rows):
                break
            data = row_map.get(q_no, {})
            co = data.get('co_marks', {})
            bl = data.get('bl_marks', {})
            total = data.get('total_marks', 2 if q_no <= 10 else (13 if q_no <= 15 else 15))

            grid.cell(row_idx, 0).text = str(q_no)
            grid.cell(row_idx, 1).text = _fmt_mark(co.get('CO1', 0))
            grid.cell(row_idx, 2).text = _fmt_mark(co.get('CO2', 0))
            grid.cell(row_idx, 3).text = _fmt_mark(co.get('CO3', 0))
            grid.cell(row_idx, 4).text = _fmt_mark(co.get('CO4', 0))
            grid.cell(row_idx, 5).text = _fmt_mark(co.get('CO5', 0))
            grid.cell(row_idx, 6).text = str(total)
            grid.cell(row_idx, 7).text = _fmt_mark(bl.get('L1', 0))
            grid.cell(row_idx, 8).text = _fmt_mark(bl.get('L2', 0))
            grid.cell(row_idx, 9).text = _fmt_mark(bl.get('L3', 0))
            grid.cell(row_idx, 10).text = _fmt_mark(bl.get('L4', 0))
            grid.cell(row_idx, 11).text = _fmt_mark(bl.get('L5', 0))
            grid.cell(row_idx, 12).text = _fmt_mark(bl.get('L6', 0))

        # Total row (r18)
        if len(grid.rows) > 18:
            co_totals = form2.get('co_totals', {})
            bl_totals = form2.get('bl_totals', {})
            grid.cell(18, 0).text = 'Total'
            grid.cell(18, 1).text = str(co_totals.get('CO1', 0))
            grid.cell(18, 2).text = str(co_totals.get('CO2', 0))
            grid.cell(18, 3).text = str(co_totals.get('CO3', 0))
            grid.cell(18, 4).text = str(co_totals.get('CO4', 0))
            grid.cell(18, 5).text = str(co_totals.get('CO5', 0))
            grid.cell(18, 6).text = str(form2.get('total_marks', 100))
            grid.cell(18, 7).text = f"L1+L2={bl_totals.get('L1', 0) + bl_totals.get('L2', 0)}"
            grid.cell(18, 9).text = f"L3+L4={bl_totals.get('L3', 0) + bl_totals.get('L4', 0)}"
            grid.cell(18, 11).text = f"L5+L6={bl_totals.get('L5', 0) + bl_totals.get('L6', 0)}"

        # Percentage row (r19)
        if len(grid.rows) > 19:
            total_marks = float(form2.get('total_marks', 0) or 0)
            co_totals = form2.get('co_totals', {})
            co_pct = {
                key: (round((float(co_totals.get(key, 0)) / total_marks) * 100) if total_marks else 0)
                for key in ['CO1', 'CO2', 'CO3', 'CO4', 'CO5']
            }
            grid.cell(19, 0).text = 'Mark Distribution in (%)'
            grid.cell(19, 1).text = f"{co_pct['CO1']}%"
            grid.cell(19, 2).text = f"{co_pct['CO2']}%"
            grid.cell(19, 3).text = f"{co_pct['CO3']}%"
            grid.cell(19, 4).text = f"{co_pct['CO4']}%"
            grid.cell(19, 5).text = f"{co_pct['CO5']}%"
            grid.cell(19, 6).text = str(form2.get('total_marks', 100))
            grid.cell(19, 7).text = f"{form2.get('l1_l2_percentage', 0)}%"
            grid.cell(19, 9).text = f"{form2.get('l3_l4_percentage', 0)}%"
            grid.cell(19, 11).text = f"{form2.get('l5_l6_percentage', 0)}%"


def _append_reference_checklist(doc, qp, form1, form2):
    """Append checklist pages from strict reference template if available."""
    template_path = _get_reference_template_path()
    if not template_path.exists():
        return False

    ref_doc = Document(str(template_path))
    _fill_reference_template(ref_doc, qp, form1, form2)

    # Ensure checklist always starts on a fresh page in the output document.
    doc.add_page_break()

    # Append only checklist pages from the reference body by locating marker paragraphs.
    source_body = ref_doc.element.body
    target_body = doc.element.body
    target_children = list(target_body)
    sect_pr = target_children[-1] if target_children and target_children[-1].tag.endswith('}sectPr') else None
    insert_pos = len(target_children) - 1 if sect_pr is not None else len(target_children)

    source_children = list(source_body)

    def _node_text(node):
        return ' '.join(t.text for t in node.iter() if t.tag.endswith('}t') and t.text)

    start_idx = None
    end_idx = None
    for idx, node in enumerate(source_children):
        txt = _node_text(node)
        if start_idx is None and 'OFFICE OF THE ADDITIONAL CONTROLLER OF EXAMINATIONS' in txt:
            start_idx = idx
        if 'Date:' in txt and 'Signature:' in txt:
            end_idx = idx

    if start_idx is None:
        return False
    if end_idx is None or end_idx < start_idx:
        end_idx = len(source_children) - 1

    for idx, element in enumerate(source_children):
        if idx < start_idx or idx > end_idx:
            continue
        target_body.insert(insert_pos, deepcopy(element))
        insert_pos += 1

    return True


def append_checklist_to_document(doc, qp):
    """Append Form 1 and Form 2 checklist sections into an existing Document.

    Returns True when checklist data was appended, False when no checklist data exists.
    """
    form1 = getattr(qp, 'submission_checklist', None) or {}
    form2 = getattr(qp, 'auto_distribution_checklist', None) or {}

    if not form1 and not form2:
        return False

    # Strict mode: use provided reference document layout exactly when available.
    if _append_reference_checklist(doc, qp, form1, form2):
        return True

    # Start a fresh section so checklist pages don't inherit unexpected layout.
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # Header block similar to the reference format.
    p = doc.add_paragraph('OFFICE OF THE ADDITIONAL CONTROLLER OF EXAMINATIONS (UDs)')
    _style_paragraph(p, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

    p = doc.add_paragraph('ANNA UNIVERSITY, CHENNAI - 600 025')
    _style_paragraph(p, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

    p = doc.add_paragraph('CHECK LIST / DECLARATION TO BE FILLED BY THE QUESTION PAPER SETTER')
    _style_paragraph(p, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, underline=True, after=10)

    details = [
        ('NAME OF THE FACULTY MEMBER', form1.get('faculty_name', '')),
        ('COURSE CODE & TITLE', form1.get('course_code_title', '')),
        ('REGULATION', form1.get('regulation', '')),
        ('MONTH & YEAR', form1.get('month_year', '')),
        ('BRANCH', form1.get('branch', '')),
    ]
    for key, value in details:
        line = doc.add_paragraph()
        line.add_run(f"{key:<33} : ").bold = True
        line.add_run(value or '-')
        _style_paragraph(line, size=11, after=2)

    checklist_items = [
        ('q1', 'Particulars regarding Regulations, Programme, Branch, Semester, Subject Code & Subject Title, Duration and Maximum Marks is clearly printed.'),
        ('q2', 'Marks for each question and / or sub-division is clearly indicated.'),
        ('q3', 'Questions are evenly distributed over all the 5 units, proportionate to the number of hours for each unit mentioned in the syllabus.'),
        ('q4', 'All the questions are within the prescribed syllabus.'),
        ('q5', 'All the figures / tables are correctly numbered and the text associated with the figures / tables are readable.'),
        ('q6', 'For each Question CO, BL are clearly specified'),
        ('q7', 'List of Tables / Charts permitted is clearly specified.'),
    ]

    for idx, (key, text) in enumerate(checklist_items, 1):
        response = _yn(form1.get(key, ''))
        para = doc.add_paragraph()
        para.add_run(f"{idx}.  {text}    -  ")
        yes_run = para.add_run('Yes')
        para.add_run(' / ')
        no_run = para.add_run('No')
        if response == 'Yes':
            yes_run.bold = True
            yes_run.underline = True
        elif response == 'No':
            no_run.bold = True
            no_run.underline = True
        _style_paragraph(para, size=11, after=2)

    p = doc.add_paragraph()
    p.add_run('NAME OF THE TABLE CHARTS: ').bold = True
    p.add_run(form1.get('table_charts_list', '') or '.................................................................')
    _style_paragraph(p, size=11, before=6, after=8)

    p = doc.add_paragraph('Recommended Distribution of Marks:')
    _style_paragraph(p, size=12, bold=True, after=4)

    rec = doc.add_table(rows=3, cols=5)
    rec.style = 'Table Grid'
    rec.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_fixed_widths(rec, [2.0, 0.6, 1.3, 1.5, 1.3])
    rec.rows[0].cells[0].text = 'Level of Questions'
    rec.rows[0].cells[1].text = ''
    rec.rows[0].cells[2].text = 'Lower Order\n(L1 and L2)'
    rec.rows[0].cells[3].text = 'Intermediate Order\n(L3 and L4)'
    rec.rows[0].cells[4].text = 'Higher Order\n(L5 and L6)'
    rec.rows[0].cells[0].merge(rec.rows[0].cells[1])

    rec.rows[1].cells[0].text = 'Recommended\nDistribution of Marks (%)'
    rec.rows[1].cells[1].text = 'UG'
    rec.rows[1].cells[2].text = '20 to 35'
    rec.rows[1].cells[3].text = 'Minimum 40'
    rec.rows[1].cells[4].text = '15 to 25'

    rec.rows[2].cells[0].text = ''
    rec.rows[2].cells[1].text = 'PG'
    rec.rows[2].cells[2].text = '10 to 25'
    rec.rows[2].cells[3].text = 'Minimum 50'
    rec.rows[2].cells[4].text = '15 to 25'
    rec.rows[1].cells[0].merge(rec.rows[2].cells[0])

    for row in rec.rows:
        for cell in row.cells:
            _style_cell(cell, size=10, bold=(row == rec.rows[0] or cell.text in ['UG', 'PG']), center=True)

    notes = [
        "Part-B Questions of 'Either OR' type should test same Bloom's Level (BL) and same Course Outcome (CO).",
        'In Parts B, Subdivisions are not compulsory and maximum subdivisions shall not exceed three.',
        'Anomalies if any in satisfying the norms of blooms taxonomy that would arise due to the choices [(Part A) 8/10 and (Part B) 4/5] may be ignored only for June/July 2020 session.',
    ]
    for note in notes:
        para = doc.add_paragraph(f'• {note}')
        _style_paragraph(para, size=10, after=1)

    doc.add_page_break()

    h2 = doc.add_paragraph()
    h2_run = h2.add_run('Checklist of Mark Distribution:')
    h2_run.bold = True
    _style_paragraph(h2, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, after=6)

    rows = form2.get('rows', [])
    if rows:
        # Two-level header to match the reference layout.
        t2 = doc.add_table(rows=2, cols=13)
        t2.style = 'Table Grid'
        t2.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_fixed_widths(t2, [0.7, 0.45, 0.45, 0.45, 0.45, 0.45, 0.55, 0.42, 0.42, 0.42, 0.42, 0.42, 0.42])

        t2.rows[0].cells[0].text = 'Question.\nNo'
        t2.rows[0].cells[1].text = 'Marks / CO'
        t2.rows[0].cells[6].text = 'Total\nMarks'
        t2.rows[0].cells[7].text = 'Marks / BL'

        t2.rows[0].cells[1].merge(t2.rows[0].cells[5])
        t2.rows[0].cells[7].merge(t2.rows[0].cells[12])
        t2.rows[0].cells[0].merge(t2.rows[1].cells[0])
        t2.rows[0].cells[6].merge(t2.rows[1].cells[6])

        sub_headers = ['CO 1', 'CO 2', 'CO 3', 'CO 4', 'CO 5', 'L1', 'L2', 'L3', 'L4', 'L5', 'L6']
        for idx, h in enumerate(sub_headers[:5], start=1):
            t2.rows[1].cells[idx].text = h
        for idx, h in enumerate(sub_headers[5:], start=7):
            t2.rows[1].cells[idx].text = h

        for row in [t2.rows[0], t2.rows[1]]:
            for cell in row.cells:
                _style_cell(cell, size=10, bold=True, center=True)

        for row_data in rows:
            row_cells = t2.add_row().cells
            row_cells[0].text = str(row_data.get('question_no', ''))
            co_marks = row_data.get('co_marks', {})
            bl_marks = row_data.get('bl_marks', {})
            row_cells[1].text = str(co_marks.get('CO1', '-') or '-')
            row_cells[2].text = str(co_marks.get('CO2', '-') or '-')
            row_cells[3].text = str(co_marks.get('CO3', '-') or '-')
            row_cells[4].text = str(co_marks.get('CO4', '-') or '-')
            row_cells[5].text = str(co_marks.get('CO5', '-') or '-')
            row_cells[6].text = str(row_data.get('total_marks', ''))
            row_cells[7].text = str(bl_marks.get('L1', '-') or '-')
            row_cells[8].text = str(bl_marks.get('L2', '-') or '-')
            row_cells[9].text = str(bl_marks.get('L3', '-') or '-')
            row_cells[10].text = str(bl_marks.get('L4', '-') or '-')
            row_cells[11].text = str(bl_marks.get('L5', '-') or '-')
            row_cells[12].text = str(bl_marks.get('L6', '-') or '-')
            for cell in row_cells:
                _style_cell(cell, size=10, center=True)

        total_row = t2.add_row().cells
        total_row[0].text = 'Total'
        co_totals = form2.get('co_totals', {})
        bl_totals = form2.get('bl_totals', {})
        total_row[1].text = str(co_totals.get('CO1', 0))
        total_row[2].text = str(co_totals.get('CO2', 0))
        total_row[3].text = str(co_totals.get('CO3', 0))
        total_row[4].text = str(co_totals.get('CO4', 0))
        total_row[5].text = str(co_totals.get('CO5', 0))
        total_row[6].text = str(form2.get('total_marks', 0))
        total_row[7].text = f"L1+L2={bl_totals.get('L1', 0) + bl_totals.get('L2', 0)}"
        total_row[8].text = ''
        total_row[9].text = f"L3+L4={bl_totals.get('L3', 0) + bl_totals.get('L4', 0)}"
        total_row[10].text = ''
        total_row[11].text = f"L5+L6={bl_totals.get('L5', 0) + bl_totals.get('L6', 0)}"
        total_row[12].text = ''
        for cell in total_row:
            _style_cell(cell, size=10, bold=True, center=True)

        pct_row = t2.add_row().cells
        pct_row[0].text = 'Mark Distribution in (%)'
        pct_row[0].merge(pct_row[5])
        pct_row[6].text = str(form2.get('total_marks', 0))
        pct_row[7].text = f"{form2.get('l1_l2_percentage', 0)}%"
        pct_row[8].text = ''
        pct_row[9].text = f"{form2.get('l3_l4_percentage', 0)}%"
        pct_row[10].text = ''
        pct_row[11].text = f"{form2.get('l5_l6_percentage', 0)}%"
        pct_row[12].text = ''
        for cell in pct_row:
            _style_cell(cell, size=10, bold=True, center=True)

        note = doc.add_paragraph('Note: In the Check list of Mark Distribution, enter the marks under corresponding Bloom\'s Level and Course Outcome (CO) in the appropriate boxes.')
        _style_paragraph(note, size=10, after=8)

        cert = doc.add_paragraph('I certify that the question paper is correct with respect to the aspects / parameters given above. The question paper may be considered for the conduct of the End - Semester Examinations.')
        _style_paragraph(cert, size=11, after=20)

        sign_line = doc.add_table(rows=1, cols=2)
        sign_line.alignment = WD_TABLE_ALIGNMENT.CENTER
        sign_line.rows[0].cells[0].text = 'Date:'
        sign_line.rows[0].cells[1].text = 'Signature:'
        _style_cell(sign_line.rows[0].cells[0], size=11)
        _style_cell(sign_line.rows[0].cells[1], size=11)
    else:
        doc.add_paragraph('Auto-filled distribution data not available.')

    return True


def build_checklist_docx_bytes(qp):
    """Build a DOCX containing Form 1 and Form 2 checklist content.

    Returns bytes or None when no checklist data exists.
    """
    doc = Document()
    appended = append_checklist_to_document(doc, qp)
    if not appended:
        return None

    buff = io.BytesIO()
    doc.save(buff)
    buff.seek(0)
    return buff.read()
